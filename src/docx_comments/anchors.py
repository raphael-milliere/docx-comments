"""Handler for comment anchors in document.xml and related story parts."""

from __future__ import annotations

import copy
from typing import TYPE_CHECKING, Iterator, Optional, Tuple

from lxml import etree

from docx_comments.xml_parts import part_element, sync_part_blob

if TYPE_CHECKING:
    from docx.document import Document
    from docx.text.paragraph import Paragraph


# OOXML Namespace
NS_W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
NS_R = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
NS_MC = "http://schemas.openxmlformats.org/markup-compatibility/2006"
NS_XML = "http://www.w3.org/XML/1998/namespace"

REL_FOOTNOTES = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes"
REL_ENDNOTES = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/endnotes"


def _qn(ns: str, name: str) -> str:
    """Create qualified name with namespace."""
    return f"{{{ns}}}{name}"


_START_TAG = _qn(NS_W, "commentRangeStart")
_END_TAG = _qn(NS_W, "commentRangeEnd")
_REF_TAG = _qn(NS_W, "commentReference")
_ID_ATTR = _qn(NS_W, "id")

# Direct children of w:p that wrap runs of visible content (EG_PContent
# members other than properties, bookmarks, proof errors, and anchors).
_RUN_CONTAINER_TAGS = frozenset(
    {
        "hyperlink",
        "ins",
        "del",
        "moveFrom",
        "moveTo",
        "smartTag",
        "sdt",
        "fldSimple",
        "dir",
        "bdo",
    }
)


def _make_reference_run(comment_id: str) -> etree._Element:
    """Build Word's styled comment reference run."""
    ref_run = etree.Element(_qn(NS_W, "r"))
    rpr = etree.SubElement(ref_run, _qn(NS_W, "rPr"))
    rstyle = etree.SubElement(rpr, _qn(NS_W, "rStyle"))
    rstyle.set(_qn(NS_W, "val"), "CommentReference")
    ref = etree.SubElement(ref_run, _REF_TAG)
    ref.set(_ID_ATTR, comment_id)
    return ref_run


class CommentAnchor:
    """Handler for comment anchors in document.xml."""

    def __init__(self, document: Document) -> None:
        self._document = document

    def _iter_anchor_parts(self) -> Iterator[Tuple[object, etree._Element]]:
        """Yield (part, XML root) pairs that can contain comment anchors."""
        seen: set[int] = set()

        def add_root(part: object, elem: Optional[etree._Element]) -> None:
            if elem is None:
                return
            elem_id = id(elem)
            if elem_id in seen:
                return
            seen.add(elem_id)
            entries.append((part, elem))

        entries: list[Tuple[object, etree._Element]] = []
        doc_part = getattr(self._document, "part", None)
        add_root(doc_part, self._document.element)

        # Headers/footers across sections, without forcing part creation.
        related_parts = getattr(doc_part, "related_parts", {}) if doc_part else {}
        for section in getattr(self._document, "sections", []):
            sect_pr = getattr(section, "_sectPr", None)
            if sect_pr is None:
                continue
            for ref_tag in ("headerReference", "footerReference"):
                for ref in sect_pr.findall(_qn(NS_W, ref_tag)):
                    r_id = ref.get(_qn(NS_R, "id"))
                    if not r_id:
                        continue
                    part = related_parts.get(r_id)
                    add_root(part, part_element(part))

        # Footnotes/endnotes parts, resolved through document relationships
        # (python-docx exposes no attribute for them and loads them as
        # generic blob Parts). External-target relationships have no part.
        if doc_part is not None:
            for rel in doc_part.rels.values():
                if rel.reltype in (REL_FOOTNOTES, REL_ENDNOTES) and not getattr(
                    rel, "is_external", False
                ):
                    part = rel.target_part
                    add_root(part, part_element(part))

        for part, root in entries:
            yield part, root

    def _iter_anchor_roots(self) -> Iterator[etree._Element]:
        """Yield XML roots that can contain comment anchors."""
        for _, root in self._iter_anchor_parts():
            yield root

    def _find_anchor_elements(
        self, comment_id: str
    ) -> tuple[
        Optional[object],
        Optional[etree._Element],
        Optional[etree._Element],
        Optional[etree._Element],
    ]:
        """Find (part, rangeStart, rangeEnd, reference) for a comment.

        Matching is done in Python (not via path predicates) so ids read from
        arbitrary documents cannot alter or break the query.

        Returns (part, None, None, reference) for comments anchored only by a
        commentReference run; (None, None, None, None) when nothing matches.
        """
        fallback: tuple = (None, None, None, None)
        for part, root in self._iter_anchor_parts():
            start = end = ref = None
            for elem in root.iter(_START_TAG, _END_TAG, _REF_TAG):
                if elem.get(_ID_ATTR) != comment_id:
                    continue
                if elem.tag == _START_TAG and start is None:
                    start = elem
                elif elem.tag == _END_TAG and end is None:
                    end = elem
                elif elem.tag == _REF_TAG and ref is None:
                    ref = elem
            if start is not None and end is not None:
                return part, start, end, ref
            if ref is not None and fallback[3] is None:
                fallback = (part, None, None, ref)

        return fallback

    def _iter_paragraphs(self) -> Iterator[Paragraph]:
        for para in self._document.paragraphs:
            yield para

        for section in getattr(self._document, "sections", []):
            for attr, ref_tag, ref_type in (
                ("header", "headerReference", None),
                ("footer", "footerReference", None),
                ("first_page_header", "headerReference", "first"),
                ("first_page_footer", "footerReference", "first"),
                ("even_page_header", "headerReference", "even"),
                ("even_page_footer", "footerReference", "even"),
            ):
                if not self._section_has_ref(section, ref_tag, ref_type):
                    continue
                part = getattr(section, attr, None)
                if part is None:
                    continue
                for para in part.paragraphs:
                    yield para

    def _section_has_ref(self, section, ref_tag: str, ref_type: Optional[str]) -> bool:
        sect_pr = getattr(section, "_sectPr", None)
        if sect_pr is None:
            return False
        for ref in sect_pr.findall(_qn(NS_W, ref_tag)):
            ref_type_attr = ref.get(_qn(NS_W, "type"))
            if ref_type is None:
                if ref_type_attr in (None, "default"):
                    return True
            elif ref_type_attr == ref_type:
                return True
        return False

    def _validate_owned(self, paragraph: Paragraph) -> None:
        """Reject paragraphs that don't live in this document's XML trees."""
        part = getattr(paragraph, "part", None)
        package = getattr(part, "package", None)
        own_package = getattr(getattr(self._document, "part", None), "package", None)
        if package is not None and own_package is not None and package is not own_package:
            raise ValueError("paragraph does not belong to this manager's document")

        # The element must live in one of this document's anchor-bearing
        # trees; a paragraph from a detached/stale tree would be mutated
        # without ever being saved.
        root = paragraph._element.getroottree().getroot()
        for _, part_root in self._iter_anchor_parts():
            if part_root is root:
                return
        raise ValueError("paragraph belongs to a detached XML tree; re-fetch it from the document")

    def _resolve_anchor_span(
        self,
        para_elem: etree._Element,
        start_run: int,
        end_run: Optional[int],
    ) -> tuple[Optional[etree._Element], Optional[etree._Element]]:
        """Resolve the (first, last) elements the anchor range should span.

        Returns (None, None) for a paragraph with no anchorable content.
        Raises IndexError/ValueError for run indices that do not address the
        paragraph's direct runs (Python-style negative indices are accepted).
        """
        runs = para_elem.findall(_qn(NS_W, "r"))
        default_span = start_run == 0 and end_run is None

        if default_span:
            # Index-free default: span all visible content in document
            # order, including runs wrapped in hyperlink/tracked-change/
            # field containers (explicit indices keep addressing direct
            # runs only, unchanged).
            anchorable = [
                child
                for child in para_elem
                # Non-element nodes (XML comments/PIs) have a non-string tag
                # and would make QName raise.
                if isinstance(child.tag, str)
                and (
                    child.tag == _qn(NS_W, "r")
                    or etree.QName(child).localname in _RUN_CONTAINER_TAGS
                )
            ]
            if anchorable:
                return anchorable[0], anchorable[-1]
            return None, None

        if not runs:
            raise IndexError(
                "paragraph has no direct runs; omit start_run/end_run to anchor the whole paragraph"
            )

        n = len(runs)
        requested_start, requested_end = start_run, end_run
        if end_run is None:
            end_run = n - 1
        if -n <= start_run < 0:
            start_run += n
        if -n <= end_run < 0:
            end_run += n
        if not 0 <= start_run < n:
            raise IndexError(
                f"start_run {requested_start} out of range for paragraph with {n} run(s)"
            )
        if not 0 <= end_run < n:
            raise IndexError(f"end_run {requested_end} out of range for paragraph with {n} run(s)")
        if end_run < start_run:
            raise ValueError(f"end_run {requested_end} precedes start_run {requested_start}")
        return runs[start_run], runs[end_run]

    def validate_anchor_target(
        self,
        paragraph: Paragraph,
        start_run: int = 0,
        end_run: Optional[int] = None,
    ) -> None:
        """
        Validate a prospective anchor target without mutating anything.

        Raises:
            ValueError: If the paragraph belongs to a different document.
            IndexError: If start_run/end_run do not address the paragraph's runs.
            ValueError: If end_run precedes start_run.
        """
        self.plan_anchor_span(paragraph, start_run, end_run)

    def plan_anchor_span(
        self,
        paragraph: Paragraph,
        start_run: int = 0,
        end_run: Optional[int] = None,
    ) -> tuple[Optional[etree._Element], Optional[etree._Element]]:
        """Validate and resolve the (first, last) elements an anchor spans.

        Resolving the span before other mutations (e.g. removing a comment's
        old anchors, which can delete reference runs and shift run indices)
        pins the range the caller actually addressed.
        """
        self._validate_owned(paragraph)
        return self._resolve_anchor_span(paragraph._element, start_run, end_run)

    def _iter_paragraph_atoms(self, para_elem):
        """Yield (child, run, length, is_text) for text atoms in order.

        Mirrors the comment-text extraction rules: w:t contributes its
        characters, w:br/w:cr/w:tab one character each; runs inside
        mc:Fallback or nested paragraphs (text boxes) are skipped.
        """
        p_tag = _qn(NS_W, "p")
        fallback = _qn(NS_MC, "Fallback")
        for run in para_elem.iter(_qn(NS_W, "r")):
            skip = False
            parent = run.getparent()
            while parent is not None and parent is not para_elem:
                if parent.tag == fallback or parent.tag == p_tag:
                    skip = True
                    break
                parent = parent.getparent()
            if skip:
                continue
            for child in run:
                # Non-element nodes (XML comments/PIs) have a non-string tag
                # and would make QName raise.
                if not isinstance(child.tag, str):
                    continue
                local = etree.QName(child).localname
                if local == "t":
                    yield child, run, len(child.text or ""), True
                elif local in ("br", "cr", "tab"):
                    yield child, run, 1, False

    def paragraph_text(self, para_elem: etree._Element) -> str:
        """The paragraph's visible text under the same rules as anchoring."""
        pieces: list[str] = []
        for child, _, _, is_text in self._iter_paragraph_atoms(para_elem):
            if is_text:
                pieces.append(child.text or "")
            else:
                local = etree.QName(child).localname
                pieces.append("\t" if local == "tab" else "\n")
        return "".join(pieces)

    def _check_char_bounds(self, para_elem: etree._Element, start_char: int, end_char: int) -> None:
        if not isinstance(start_char, int) or not isinstance(end_char, int):
            raise TypeError("start_char and end_char must be integers")
        if start_char < 0 or end_char < 0:
            raise IndexError("character offsets must be non-negative")
        if end_char <= start_char:
            raise ValueError(f"end_char {end_char} must be greater than start_char {start_char}")
        total = sum(length for _, _, length, _ in self._iter_paragraph_atoms(para_elem))
        if end_char > total:
            raise IndexError(
                f"end_char {end_char} out of range for paragraph with {total} character(s)"
            )

    def validate_char_span(self, paragraph: Paragraph, start_char: int, end_char: int) -> None:
        """Validate a character span without mutating anything."""
        self._validate_owned(paragraph)
        self._check_char_bounds(paragraph._element, start_char, end_char)

    def _split_run_at_child(self, run, child):
        """Split `run` immediately before `child`. Returns (left, right)."""
        left = etree.Element(_qn(NS_W, "r"))
        for key, value in run.attrib.items():
            left.set(key, value)
        rpr = run.find(_qn(NS_W, "rPr"))
        if rpr is not None:
            left.append(copy.deepcopy(rpr))
        for sibling in list(run):
            if sibling is child:
                break
            if sibling.tag == _qn(NS_W, "rPr"):
                continue
            left.append(sibling)  # moves the element out of `run`
        run.addprevious(left)
        return left, run

    def _split_run_at_text(self, run, t_child, offset):
        """Split a run inside its w:t at `offset` characters. Returns (left, right)."""
        text = t_child.text or ""
        left_run, right_run = self._split_run_at_child(run, t_child)
        left_t = etree.SubElement(left_run, _qn(NS_W, "t"))
        left_text = text[:offset]
        left_t.text = left_text
        if left_text.strip() != left_text:
            left_t.set(_qn(NS_XML, "space"), "preserve")
        right_text = text[offset:]
        t_child.text = right_text
        if right_text.strip() != right_text:
            t_child.set(_qn(NS_XML, "space"), "preserve")
        else:
            t_child.attrib.pop(_qn(NS_XML, "space"), None)
        return left_run, right_run

    def _ensure_run_boundary(self, para_elem, pos):
        """Make a run boundary exist at char offset `pos`.

        Returns (run_ending_at_pos, run_starting_at_pos); either side is
        None at the paragraph edges.
        """
        cum = 0
        prev_run = None
        for child, run, length, is_text in self._iter_paragraph_atoms(para_elem):
            if cum == pos:
                if prev_run is run:
                    return self._split_run_at_child(run, child)
                return prev_run, run
            if cum < pos < cum + length:
                return self._split_run_at_text(run, child, pos - cum)
            cum += length
            prev_run = run
        return prev_run, None

    def add_anchors_at_char_span(
        self,
        paragraph: Paragraph,
        start_char: int,
        end_char: int,
        comment_id: str,
    ) -> None:
        """Anchor a comment to an exact character span, splitting runs.

        Splitting preserves formatting (rPr is deep-copied) and whitespace
        (xml:space="preserve" on chunks with edge whitespace); the
        paragraph's visible text is unchanged.
        """
        self._validate_owned(paragraph)
        para_elem = paragraph._element
        self._check_char_bounds(para_elem, start_char, end_char)
        # End boundary first: when both offsets fall inside one run, the
        # start split keeps the original element as the right half (still
        # ending at end_char), so last_run stays valid; splitting start
        # first would let the end split re-split first_run and leave it
        # pointing at the tail after end_char.
        last_run, _ = self._ensure_run_boundary(para_elem, end_char)
        _, first_run = self._ensure_run_boundary(para_elem, start_char)
        self.add_anchors_at_span(paragraph, (first_run, last_run), comment_id)

    def _sync_containing_part(self, elem: etree._Element) -> None:
        """Persist the blob-backed part (if any) that owns this element."""
        root = elem.getroottree().getroot()
        for part, part_root in self._iter_anchor_parts():
            if part_root is root:
                sync_part_blob(part)
                return

    @staticmethod
    def ensure_span_survives_removal(
        span: tuple[Optional[etree._Element], Optional[etree._Element]],
        comment_ids: set[str],
    ) -> None:
        """Raise before any mutation when a span endpoint would be deleted.

        remove_anchors() deletes a reference run only when nothing but run
        properties and matching references would remain — the same condition
        is mirrored here so runs that also carry text are not rejected.
        """
        for endpoint in span:
            if endpoint is None or etree.QName(endpoint).localname != "r":
                continue
            matching = [
                child
                for child in endpoint
                if child.tag == _REF_TAG and child.get(_ID_ATTR) in comment_ids
            ]
            if matching and all(
                child in matching or etree.QName(child).localname == "rPr" for child in endpoint
            ):
                raise ValueError(
                    "start_run/end_run address a comment reference run "
                    "that this operation removes; recompute the indices"
                )

    def add_anchors_at_span(
        self,
        paragraph: Paragraph,
        span: tuple[Optional[etree._Element], Optional[etree._Element]],
        comment_id: str,
    ) -> None:
        """Insert anchors around a span previously resolved by plan_anchor_span.

        Raises:
            ValueError: If a span endpoint no longer exists in the paragraph
                (e.g. the indices addressed a comment reference run that a
                preceding remove_anchors() deleted).
        """
        first, last = span
        para_elem = paragraph._element

        if first is None or last is None:
            # No anchorable content: anchor at paragraph level.
            self._add_anchors_to_empty_paragraph(para_elem, comment_id)
            self._sync_containing_part(para_elem)
            return

        if first.getparent() is None or last.getparent() is None:
            raise ValueError(
                "anchor target runs no longer exist (start_run/end_run "
                "addressed a comment reference run removed by this "
                "operation); recompute the indices"
            )

        # Insert commentRangeStart before the first spanned element
        range_start = etree.Element(_START_TAG)
        range_start.set(_ID_ATTR, comment_id)
        first.addprevious(range_start)

        # Insert commentRangeEnd after the last spanned element
        range_end = etree.Element(_END_TAG)
        range_end.set(_ID_ATTR, comment_id)
        last.addnext(range_end)

        # Insert commentReference run after commentRangeEnd
        ref_run = _make_reference_run(comment_id)
        range_end.addnext(ref_run)

        # Persist for blob-backed parts (footnotes/endnotes).
        self._sync_containing_part(para_elem)

    def add_anchors(
        self,
        paragraph: Paragraph,
        comment_id: str,
        start_run: int = 0,
        end_run: Optional[int] = None,
    ) -> None:
        """
        Add comment anchors to a paragraph.

        Creates commentRangeStart, commentRangeEnd, and commentReference
        elements around the specified runs. Runs are the paragraph's direct
        w:r children; Python-style negative indices are accepted, and invalid
        indices raise (IndexError/ValueError) rather than being adjusted.

        Args:
            paragraph: The paragraph to anchor the comment to.
            comment_id: The comment ID.
            start_run: Index of first run to anchor (default: 0).
            end_run: Index of last run to anchor (default: last run).

        Raises:
            ValueError: If the paragraph belongs to a different document, or
                end_run precedes start_run.
            IndexError: If start_run/end_run are out of range.
        """
        span = self.plan_anchor_span(paragraph, start_run, end_run)
        self.add_anchors_at_span(paragraph, span, comment_id)

    def _add_anchors_to_empty_paragraph(
        self,
        para_elem: etree._Element,
        comment_id: str,
    ) -> None:
        """Add anchors to a paragraph with no runs."""
        # Create commentRangeStart
        range_start = etree.Element(_START_TAG)
        range_start.set(_ID_ATTR, comment_id)

        # Create commentRangeEnd
        range_end = etree.Element(_END_TAG)
        range_end.set(_ID_ATTR, comment_id)

        # Create commentReference run
        ref_run = _make_reference_run(comment_id)

        # Insert after pPr if present, else at start
        pPr = para_elem.find(_qn(NS_W, "pPr"))
        if pPr is not None:
            pPr.addnext(range_start)
        else:
            para_elem.insert(0, range_start)

        range_start.addnext(range_end)
        range_end.addnext(ref_run)

    @staticmethod
    def _has_paragraph_ancestor(elem: etree._Element) -> bool:
        parent = elem.getparent()
        while parent is not None:
            if etree.QName(parent).localname == "p":
                return True
            parent = parent.getparent()
        return False

    @staticmethod
    def _last_paragraph_before(elem: etree._Element) -> Optional[etree._Element]:
        """Last w:p in document order before `elem`, searching up the tree."""
        p_tag = _qn(NS_W, "p")
        current: Optional[etree._Element] = elem
        while current is not None:
            for sib in current.itersiblings(preceding=True):
                if sib.tag == p_tag:
                    return sib
                inner = sib.findall(f".//{p_tag}")
                if inner:
                    return inner[-1]
            current = current.getparent()
            if current is not None and current.tag == p_tag:
                return current
        return None

    def _add_anchors_around_reference(
        self, parent_ref: etree._Element, new_comment_id: str
    ) -> None:
        """Anchor a reply around a parent that has only a reference run."""
        ref_run_parent = parent_ref.getparent()
        target = (
            ref_run_parent
            if ref_run_parent is not None and etree.QName(ref_run_parent).localname == "r"
            else parent_ref
        )
        new_start = etree.Element(_START_TAG)
        new_start.set(_ID_ATTR, new_comment_id)
        target.addprevious(new_start)
        new_end = etree.Element(_END_TAG)
        new_end.set(_ID_ATTR, new_comment_id)
        target.addnext(new_end)
        new_end.addnext(_make_reference_run(new_comment_id))

    def add_anchors_at_comment(
        self,
        parent_comment_id: str,
        new_comment_id: str,
    ) -> None:
        """
        Add anchors for a new comment at the same location as an existing comment.

        Used for reply comments that should anchor to the same text.

        Args:
            parent_comment_id: ID of the existing comment.
            new_comment_id: ID of the new comment.
        """
        # Find the parent comment's anchors
        part, parent_start, parent_end, parent_ref = self._find_anchor_elements(parent_comment_id)

        if parent_start is None or parent_end is None:
            if parent_ref is None:
                raise ValueError(f"Could not find anchors for comment {parent_comment_id}")
            # Reference-only anchor (range markers are optional per ECMA-376
            # §17.13.4): synthesize a range around the parent's reference run
            # for the new comment, mirroring what Word produces on re-save.
            self._add_anchors_around_reference(parent_ref, new_comment_id)
            sync_part_blob(part)
            return

        # Add new anchors after any existing anchor group for this location.
        def is_comment_ref_run(elem: etree._Element) -> bool:
            if etree.QName(elem).localname != "r":
                return False
            return elem.find(_REF_TAG) is not None

        # Insert new start after the last commentRangeStart in the group.
        insert_start_after = parent_start
        sibling = parent_start.getnext()
        while sibling is not None and etree.QName(sibling).localname == "commentRangeStart":
            insert_start_after = sibling
            sibling = sibling.getnext()

        new_start = etree.Element(_START_TAG)
        new_start.set(_ID_ATTR, new_comment_id)
        insert_start_after.addnext(new_start)

        # Insert new end after the last commentRangeEnd in the group.
        insert_end_after = parent_end
        sibling = parent_end.getnext()
        while sibling is not None and etree.QName(sibling).localname == "commentRangeEnd":
            insert_end_after = sibling
            sibling = sibling.getnext()

        new_end = etree.Element(_END_TAG)
        new_end.set(_ID_ATTR, new_comment_id)
        insert_end_after.addnext(new_end)

        # Place the reference run at a schema-valid run position. A bare run
        # directly under w:body/w:tbl/w:tr/w:tc (block-level ranges) makes
        # document.xml invalid and triggers Word's repair prompt.
        ref_run = _make_reference_run(new_comment_id)

        anchor_after: Optional[etree._Element] = None
        if parent_ref is not None:
            parent_ref_run = parent_ref.getparent()
            if parent_ref_run is not None and etree.QName(parent_ref_run).localname == "r":
                anchor_after = parent_ref_run
        if anchor_after is None and self._has_paragraph_ancestor(new_end):
            anchor_after = new_end

        if anchor_after is not None:
            insert_ref_after = anchor_after
            sibling = insert_ref_after.getnext()
            while sibling is not None and is_comment_ref_run(sibling):
                insert_ref_after = sibling
                sibling = sibling.getnext()
            insert_ref_after.addnext(ref_run)
        else:
            para = self._last_paragraph_before(new_end)
            if para is None:
                raise ValueError(
                    f"cannot place the comment reference run for comment "
                    f"{new_comment_id}: no paragraph found within the range "
                    f"of comment {parent_comment_id}"
                )
            para.append(ref_run)

        # Persist for blob-backed parts (footnotes/endnotes).
        sync_part_blob(part)

    def find_paragraph_with_comment(self, comment_id: str) -> Optional[Paragraph]:
        """
        Find the paragraph that contains a comment's anchor.

        Args:
            comment_id: The comment ID to find.

        Returns:
            The Paragraph object, or None if not found. Paragraphs in
            footnotes/endnotes are returned as element-wrapping proxies whose
            .part is the main document part; treat them as read-mostly (edits
            made through such a proxy are not synced to the footnotes part).
        """
        # Find commentRangeStart for this comment
        _, range_start, _, _ = self._find_anchor_elements(comment_id)

        if range_start is None:
            return None

        # Walk up to find parent paragraph
        parent = range_start.getparent()
        while parent is not None:
            if etree.QName(parent).localname == "p":
                # Find matching python-docx Paragraph
                for para in self._iter_paragraphs():
                    if para._element is parent:
                        return para
                # Paragraphs inside tables, text boxes, footnotes, etc. are
                # not reachable through document.paragraphs; wrap the element
                # directly instead of reporting the comment as unanchored.
                from docx.text.paragraph import Paragraph as _Paragraph

                return _Paragraph(parent, self._document)
            parent = parent.getparent()

        return None

    @staticmethod
    def _in_fallback(elem: etree._Element) -> bool:
        parent = elem.getparent()
        fallback = _qn(NS_MC, "Fallback")
        while parent is not None:
            if parent.tag == fallback:
                return True
            parent = parent.getparent()
        return False

    def get_anchored_text(self, comment_id: str) -> Optional[str]:
        """Text between a comment's range markers, in document order.

        Mirrors the extraction rules of comment-text reading: w:t text,
        w:br/w:cr as newline, w:tab as tab, mc:Fallback content skipped,
        and a newline when the range crosses a paragraph boundary.
        Returns None when the comment has no commentRangeStart/End pair.
        """
        _, start, end, _ = self._find_anchor_elements(comment_id)
        if start is None or end is None:
            return None
        root = start.getroottree().getroot()
        p_tag = _qn(NS_W, "p")
        r_tag = _qn(NS_W, "r")
        start_in_paragraph = self._has_paragraph_ancestor(start)
        pieces: list[str] = []
        active = False
        paras_seen = 0
        for elem in root.iter():
            if elem is start:
                active = True
                continue
            if elem is end:
                break
            if not active:
                continue
            if elem.tag == p_tag:
                if paras_seen > 0 or start_in_paragraph:
                    pieces.append("\n")
                paras_seen += 1
                continue
            parent = elem.getparent()
            if parent is None or parent.tag != r_tag:
                continue  # e.g. w:tab inside w:pPr/w:tabs
            if self._in_fallback(elem):
                continue
            local = etree.QName(elem).localname
            if local == "t":
                if elem.text:
                    pieces.append(elem.text)
            elif local in ("br", "cr"):
                pieces.append("\n")
            elif local == "tab":
                pieces.append("\t")
        return "".join(pieces)

    def remove_anchors(self, comment_id: str) -> None:
        """
        Remove all anchors for a comment.

        Args:
            comment_id: The comment ID whose anchors to remove.
        """
        for part, root in self._iter_anchor_parts():
            mutated = False

            for tag in (_START_TAG, _END_TAG):
                for elem in list(root.iter(tag)):
                    if elem.get(_ID_ATTR) != comment_id:
                        continue
                    elem.getparent().remove(elem)
                    mutated = True

            # Find and remove commentReference (and its parent run when only
            # run properties would remain, e.g. Word's rPr/rStyle wrapper).
            for ref in list(root.iter(_REF_TAG)):
                if ref.get(_ID_ATTR) != comment_id:
                    continue
                ref_run = ref.getparent()
                if (
                    ref_run is not None
                    and etree.QName(ref_run).localname == "r"
                    and all(
                        child is ref or etree.QName(child).localname == "rPr" for child in ref_run
                    )
                ):
                    ref_run.getparent().remove(ref_run)
                else:
                    ref.getparent().remove(ref)
                mutated = True

            if mutated:
                sync_part_blob(part)

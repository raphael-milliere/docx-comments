"""Timestamp control and date parsing."""

import zipfile
from datetime import datetime, timezone

from docx import Document
from lxml import etree

from docx_comments import CommentManager, PersonInfo
from docx_comments.manager import _parse_comment_date

NS_W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
NS_W16CEX = "http://schemas.microsoft.com/office/word/2018/wordml/cex"


class TestParseCommentDate:
    def test_z_form(self):
        parsed = _parse_comment_date("2020-01-02T03:04:05Z")
        assert parsed == datetime(2020, 1, 2, 3, 4, 5, tzinfo=timezone.utc)

    def test_offset_form(self):
        parsed = _parse_comment_date("2020-01-02T05:04:05+02:00")
        assert parsed == datetime(2020, 1, 2, 3, 4, 5, tzinfo=timezone.utc)

    def test_naive_assumed_utc(self):
        parsed = _parse_comment_date("2020-01-02T03:04:05")
        assert parsed.tzinfo == timezone.utc

    def test_garbage_returns_none(self):
        assert _parse_comment_date("not-a-date") is None

    def test_empty_returns_none(self):
        assert _parse_comment_date(None) is None
        assert _parse_comment_date("") is None


class TestTimestampParameter:
    def test_round_trip_and_date_utc(self, tmp_path):
        doc = Document()
        para = doc.add_paragraph("text")
        mgr = CommentManager(doc)
        stamp = datetime(2020, 1, 2, 3, 4, 5, tzinfo=timezone.utc)
        mgr.add_comment(para, "c", PersonInfo(author="A"), timestamp=stamp)
        path = tmp_path / "t.docx"
        doc.save(str(path))
        doc2 = Document(str(path))
        comment = next(iter(CommentManager(doc2).list_comments()))
        assert comment.timestamp == stamp
        with zipfile.ZipFile(path) as zf:
            ext = etree.fromstring(zf.read("word/commentsExtensible.xml"))
        entries = [
            e.get(f"{{{NS_W16CEX}}}dateUtc")
            for e in ext
            if etree.QName(e).localname == "commentExtensible"
        ]
        assert entries == ["2020-01-02T03:04:05Z"]

    def test_naive_timestamp_treated_as_local(self):
        doc = Document()
        para = doc.add_paragraph("text")
        mgr = CommentManager(doc)
        naive = datetime(2020, 6, 1, 12, 0, 0)
        mgr.add_comment(para, "c", PersonInfo(author="A"), timestamp=naive)
        comment = next(iter(mgr.list_comments()))
        assert comment.timestamp == naive.astimezone(timezone.utc)

    def test_reply_timestamp_orders_replies(self):
        doc = Document()
        para = doc.add_paragraph("text")
        mgr = CommentManager(doc)
        cid = mgr.add_comment(para, "root", PersonInfo(author="A"))
        t1 = datetime(2020, 1, 1, tzinfo=timezone.utc)
        t2 = datetime(2020, 1, 2, tzinfo=timezone.utc)
        r_late = mgr.reply_to_comment(cid, "late", PersonInfo(author="B"), timestamp=t2)
        r_early = mgr.reply_to_comment(
            cid, "early", PersonInfo(author="B"), timestamp=t1
        )
        thread = mgr.get_comment_threads()[0]
        assert [r.comment_id for r in thread.replies] == [r_early, r_late]

    def test_threads_survive_corrupted_date(self):
        doc = Document()
        para = doc.add_paragraph("text")
        mgr = CommentManager(doc)
        mgr.add_comment(para, "c", PersonInfo(author="A"))
        for elem in mgr._comments_xml.findall(f"{{{NS_W}}}comment"):
            elem.set(f"{{{NS_W}}}date", "garbage")
        threads = mgr.get_comment_threads()
        assert len(threads) == 1
        assert threads[0].root.timestamp is None

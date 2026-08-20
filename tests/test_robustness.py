"""Regression tests for defects found in the 2026-08 adversarial review."""

from zipfile import ZipFile

import pytest
from docx import Document
from lxml import etree

from docx_comments import CommentManager, CommentNotFoundError, PersonInfo
from docx_comments.anchors import REL_FOOTNOTES, CommentAnchor
from docx_comments.xml_parts import (
    CommentsExtendedPart,
    CommentsIdsPart,
    CommentsPart,
    parse_xml_bytes,
    part_element,
)

NS_W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
NS_W14 = "http://schemas.microsoft.com/office/word/2010/wordml"
NS_W15 = "http://schemas.microsoft.com/office/word/2012/wordml"
NS_XML = "http://www.w3.org/XML/1998/namespace"
NS_MC = "http://schemas.openxmlformats.org/markup-compatibility/2006"

CT_FOOTNOTES = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.footnotes+xml"
)

HAS_NATIVE_COMMENTS = hasattr(Document(), "add_comment")


def qn(ns: str, name: str) -> str:
    return f"{{{ns}}}{name}"


def author_obj(name: str) -> PersonInfo:
    return PersonInfo(author=name)


def saved_zip_names(doc, tmp_path, name="probe.docx"):
    path = tmp_path / name
    doc.save(str(path))
    with ZipFile(str(path)) as zf:
        return set(zf.namelist()), path


class TestIdGeneration:
    """Generated ids must be spec-valid and unique."""

    def test_comment_ids_within_int32(self, tmp_path):
        doc = Document()
        mgr = CommentManager(doc)
        for i in range(10):
            para = doc.add_paragraph(f"Paragraph {i}")
            mgr.add_comment(para, f"Comment {i}", author_obj("A"))

        _, path = saved_zip_names(doc, tmp_path)
        with ZipFile(str(path)) as zf:
            comments = etree.fromstring(zf.read("word/comments.xml"))
            document = etree.fromstring(zf.read("word/document.xml"))

        ids = [c.get(qn(NS_W, "id")) for c in comments.findall(qn(NS_W, "comment"))]
        for tag in ("commentRangeStart", "commentRangeEnd", "commentReference"):
            ids.extend(e.get(qn(NS_W, "id")) for e in document.iter(qn(NS_W, tag)))

        assert len(ids) >= 30
        for value in ids:
            assert 0 < int(value) <= 0x7FFFFFFE

        # paraId / textId are ST_LongHexNumber within the valid range
        for para in comments.iter(qn(NS_W, "p")):
            for attr in ("paraId", "textId"):
                value = para.get(qn(NS_W14, attr))
                assert value is not None
                assert 0 < int(value, 16) <= 0x7FFFFFFE

    def test_global_random_seed_does_not_collide_ids(self):
        import random

        doc = Document()
        mgr = CommentManager(doc)
        random.seed(42)
        id1 = mgr.add_comment(doc.add_paragraph("One"), "c1", author_obj("A"))
        random.seed(42)
        id2 = mgr.add_comment(doc.add_paragraph("Two"), "c2", author_obj("A"))
        assert id1 != id2
        para_ids = {c.para_id for c in mgr.list_comments()}
        assert len(para_ids) == 2

    def test_comment_id_collision_redrawn(self, monkeypatch):
        from docx_comments import manager as manager_mod

        seq = iter(["123", "123", "456"])
        monkeypatch.setattr(manager_mod, "_generate_id", lambda: next(seq))

        doc = Document()
        mgr = CommentManager(doc)
        id1 = mgr.add_comment(doc.add_paragraph("One"), "c1", author_obj("A"))
        id2 = mgr.add_comment(doc.add_paragraph("Two"), "c2", author_obj("A"))
        assert id1 == "123"
        assert id2 == "456"

    def test_hex_id_collision_redrawn(self, monkeypatch):
        from docx_comments import manager as manager_mod

        values = ["AAAAAAA1", "AAAAAAA1", "AAAAAAA2", "AAAAAAA3", "AAAAAAA4"]
        seq = iter(values)
        monkeypatch.setattr(manager_mod, "_generate_long_hex_id", lambda: next(seq))

        doc = Document()
        mgr = CommentManager(doc)
        mgr.add_comment(doc.add_paragraph("One"), "c1", author_obj("A"))
        comment = next(iter(mgr.list_comments()))
        # para_id, text_id, durable_id must all be distinct despite the
        # colliding generator draw.
        assert comment.para_id == "AAAAAAA1"
        assert comment.durable_id in ("AAAAAAA3", "AAAAAAA4")


class TestResolutionSemantics:
    """Resolution is thread-scoped, like Word."""

    def test_resolve_marks_whole_thread(self):
        doc = Document()
        para = doc.add_paragraph("Text")
        mgr = CommentManager(doc)
        root_id = mgr.add_comment(para, "Root", author_obj("A"))
        mgr.reply_to_comment(root_id, "Reply 1", author_obj("B"))
        reply2_id = mgr.reply_to_comment(root_id, "Reply 2", author_obj("C"))

        mgr.resolve_comment(root_id)
        threading = CommentsExtendedPart(doc).get_threading_info()
        assert len(threading) == 3
        assert all(entry["done"] for entry in threading.values())
        assert mgr.get_comment_threads()[0].is_resolved

        # Unresolving via a reply reopens the whole thread.
        mgr.unresolve_comment(reply2_id)
        threading = CommentsExtendedPart(doc).get_threading_info()
        assert not any(entry["done"] for entry in threading.values())

    def test_resolve_error_names_the_comment_id(self):
        doc = Document()
        mgr = CommentManager(doc)
        with pytest.raises(ValueError, match="no-such-id"):
            mgr.resolve_comment("no-such-id")

    @pytest.mark.skipif(not HAS_NATIVE_COMMENTS, reason="needs python-docx >= 1.2")
    def test_resolve_native_python_docx_comment(self):
        doc = Document()
        para = doc.add_paragraph("Native text")
        doc.add_comment(runs=para.runs, text="native", author="Nat", initials="N")

        mgr = CommentManager(doc)
        comments = list(mgr.list_comments())
        assert len(comments) == 1
        mgr.resolve_comment(comments[0].comment_id)
        assert next(iter(mgr.list_comments())).is_resolved


class TestTextFidelity:
    """Comment text must survive the write/read round trip."""

    def test_whitespace_preserved(self, tmp_path):
        doc = Document()
        mgr = CommentManager(doc)
        mgr.add_comment(doc.add_paragraph("Text"), "  padded  ", author_obj("A"))

        _, path = saved_zip_names(doc, tmp_path)
        with ZipFile(str(path)) as zf:
            comments = etree.fromstring(zf.read("word/comments.xml"))
        t_elems = [t for t in comments.iter(qn(NS_W, "t")) if t.text]
        assert t_elems[0].text == "  padded  "
        assert t_elems[0].get(qn(NS_XML, "space")) == "preserve"

        mgr2 = CommentManager(Document(str(path)))
        assert next(iter(mgr2.list_comments())).text == "  padded  "

    def test_newlines_and_tabs_roundtrip(self, tmp_path):
        text = "line1\nline2\ttabbed"
        doc = Document()
        mgr = CommentManager(doc)
        mgr.add_comment(doc.add_paragraph("Text"), text, author_obj("A"))

        _, path = saved_zip_names(doc, tmp_path)
        with ZipFile(str(path)) as zf:
            comments = etree.fromstring(zf.read("word/comments.xml"))
        assert comments.find(f".//{qn(NS_W, 'br')}") is not None
        assert comments.find(f".//{qn(NS_W, 'tab')}") is not None
        for t in comments.iter(qn(NS_W, "t")):
            assert "\n" not in (t.text or "")
            assert "\t" not in (t.text or "")

        mgr2 = CommentManager(Document(str(path)))
        assert next(iter(mgr2.list_comments())).text == text

    def test_empty_text_roundtrip(self, tmp_path):
        doc = Document()
        mgr = CommentManager(doc)
        mgr.add_comment(doc.add_paragraph("Text"), "", author_obj("A"))
        _, path = saved_zip_names(doc, tmp_path)
        mgr2 = CommentManager(Document(str(path)))
        assert next(iter(mgr2.list_comments())).text == ""

    @pytest.mark.skipif(not HAS_NATIVE_COMMENTS, reason="needs python-docx >= 1.2")
    def test_multiparagraph_native_comment_readback(self):
        doc = Document()
        para = doc.add_paragraph("Native text")
        doc.add_comment(runs=para.runs, text="para1\npara2", author="Nat")

        mgr = CommentManager(doc)
        assert next(iter(mgr.list_comments())).text == "para1\npara2"

    def test_illegal_control_chars_raise_without_mutation(self, tmp_path):
        doc = Document()
        para = doc.add_paragraph("Text")
        mgr = CommentManager(doc)
        with pytest.raises(ValueError, match="not allowed in XML"):
            mgr.add_comment(para, "bad\x00text", author_obj("A"))

        names, _ = saved_zip_names(doc, tmp_path)
        assert "word/comments.xml" not in names
        assert para._element.find(qn(NS_W, "commentRangeStart")) is None


class TestAnchorValidation:
    """Run indices are validated, not silently rewritten."""

    def test_out_of_range_end_run_raises(self):
        doc = Document()
        para = doc.add_paragraph("Single run")
        mgr = CommentManager(doc)
        with pytest.raises(IndexError):
            mgr.add_comment(para, "c", author_obj("A"), end_run=2)

    def test_out_of_range_start_run_raises(self):
        doc = Document()
        para = doc.add_paragraph("Single run")
        mgr = CommentManager(doc)
        with pytest.raises(IndexError):
            mgr.add_comment(para, "c", author_obj("A"), start_run=5)

    def test_inverted_range_raises(self):
        doc = Document()
        para = doc.add_paragraph("")
        for word in ("one ", "two ", "three"):
            para.add_run(word)
        mgr = CommentManager(doc)
        with pytest.raises(ValueError):
            mgr.add_comment(para, "c", author_obj("A"), start_run=2, end_run=1)

    def test_negative_indices_address_from_the_end(self):
        doc = Document()
        para = doc.add_paragraph("")
        for word in ("one ", "two ", "three"):
            para.add_run(word)
        mgr = CommentManager(doc)
        comment_id = mgr.add_comment(
            para, "c", author_obj("A"), start_run=-1, end_run=-1
        )

        children = list(para._element)
        localnames = [etree.QName(c).localname for c in children]
        start_idx = localnames.index("commentRangeStart")
        # The last run is wrapped: rangeStart sits after the first two runs.
        assert localnames[:start_idx].count("r") == 2
        assert children[start_idx].get(qn(NS_W, "id")) == comment_id

    def test_cross_document_paragraph_raises(self, tmp_path):
        doc_a = Document()
        doc_b = Document()
        doc_b_para = doc_b.add_paragraph("Other document")
        mgr = CommentManager(doc_a)
        with pytest.raises(ValueError, match="does not belong"):
            mgr.add_comment(doc_b_para, "c", author_obj("A"))

        names_a, _ = saved_zip_names(doc_a, tmp_path, "a.docx")
        assert "word/comments.xml" not in names_a
        assert doc_b_para._element.find(qn(NS_W, "commentRangeStart")) is None

    def test_hyperlink_only_paragraph_anchors_whole_content(self):
        doc = Document()
        para = doc.add_paragraph("")
        hyperlink = etree.SubElement(para._element, qn(NS_W, "hyperlink"))
        run = etree.SubElement(hyperlink, qn(NS_W, "r"))
        t = etree.SubElement(run, qn(NS_W, "t"))
        t.text = "link text"

        mgr = CommentManager(doc)
        mgr.add_comment(para, "on a link", author_obj("A"))

        localnames = [etree.QName(c).localname for c in para._element]
        assert localnames.index("commentRangeStart") < localnames.index("hyperlink")
        assert localnames.index("hyperlink") < localnames.index("commentRangeEnd")


def _append_hyperlink(para, text):
    hyperlink = etree.SubElement(para._element, qn(NS_W, "hyperlink"))
    run = etree.SubElement(hyperlink, qn(NS_W, "r"))
    t = etree.SubElement(run, qn(NS_W, "t"))
    t.text = text
    return hyperlink


class TestMixedContentAnchors:
    def _child_locals(self, para):
        return [etree.QName(c).localname for c in para._element]

    def test_run_then_hyperlink_default_anchor_covers_all(self):
        doc = Document()
        para = doc.add_paragraph()
        para.add_run("See ")
        _append_hyperlink(para, "the link")
        mgr = CommentManager(doc)
        mgr.add_comment(para, "c", PersonInfo(author="A"))
        locals_ = self._child_locals(para)
        assert locals_.index("commentRangeEnd") > locals_.index("hyperlink"), (
            f"anchor truncated before the hyperlink: {locals_}"
        )

    def test_hyperlink_then_run_default_anchor_covers_all(self):
        doc = Document()
        para = doc.add_paragraph()
        _append_hyperlink(para, "the link")
        para.add_run(" trailing")
        mgr = CommentManager(doc)
        mgr.add_comment(para, "c", PersonInfo(author="A"))
        locals_ = self._child_locals(para)
        assert locals_.index("commentRangeStart") < locals_.index("hyperlink")

    def test_tracked_change_wrapper_covered(self):
        doc = Document()
        para = doc.add_paragraph()
        ins = etree.SubElement(para._element, qn(NS_W, "ins"))
        run = etree.SubElement(ins, qn(NS_W, "r"))
        t = etree.SubElement(run, qn(NS_W, "t"))
        t.text = "inserted"
        para.add_run(" kept")
        mgr = CommentManager(doc)
        mgr.add_comment(para, "c", PersonInfo(author="A"))
        locals_ = self._child_locals(para)
        assert locals_.index("commentRangeStart") < locals_.index("ins")

    def test_xml_comment_node_in_paragraph_is_skipped(self):
        doc = Document()
        para = doc.add_paragraph()
        para._element.append(etree.Comment("note"))
        para.add_run("text")
        mgr = CommentManager(doc)
        mgr.add_comment(para, "c", PersonInfo(author="A"))
        children = list(para._element)
        start_idx = next(
            i for i, c in enumerate(children) if c.tag == qn(NS_W, "commentRangeStart")
        )
        end_idx = next(
            i for i, c in enumerate(children) if c.tag == qn(NS_W, "commentRangeEnd")
        )
        run_idx = next(
            i
            for i, c in enumerate(children)
            if c.tag == qn(NS_W, "r") and c.find(qn(NS_W, "t")) is not None
        )
        assert start_idx < run_idx < end_idx

    def test_explicit_indices_still_address_direct_runs(self):
        doc = Document()
        para = doc.add_paragraph()
        para.add_run("one")
        _append_hyperlink(para, "link")
        para.add_run("two")
        mgr = CommentManager(doc)
        mgr.add_comment(para, "c", PersonInfo(author="A"), start_run=0, end_run=0)
        children = list(para._element)
        end_idx = next(
            i for i, c in enumerate(children)
            if etree.QName(c).localname == "commentRangeEnd"
        )
        hyper_idx = next(
            i for i, c in enumerate(children)
            if etree.QName(c).localname == "hyperlink"
        )
        assert end_idx < hyper_idx, "explicit run indices must not span containers"

    def test_reference_run_carries_comment_reference_style(self):
        doc = Document()
        para = doc.add_paragraph("text")
        mgr = CommentManager(doc)
        mgr.add_comment(para, "c", PersonInfo(author="A"))
        ref = para._element.find(f".//{qn(NS_W, 'commentReference')}")
        ref_run = ref.getparent()
        style = ref_run.find(f"{qn(NS_W, 'rPr')}/{qn(NS_W, 'rStyle')}")
        assert style is not None and style.get(qn(NS_W, "val")) == "CommentReference"

    def test_styled_reference_run_still_removed_on_delete(self):
        doc = Document()
        para = doc.add_paragraph("text")
        mgr = CommentManager(doc)
        cid = mgr.add_comment(para, "c", PersonInfo(author="A"))
        mgr.delete_comment(cid)
        assert para._element.find(f".//{qn(NS_W, 'commentReference')}") is None
        assert para._element.find(qn(NS_W, "r")) is not None  # text run kept


class TestAnchorRemoval:
    def test_word_style_reference_run_fully_removed(self):
        doc = Document()
        para = doc.add_paragraph("")
        for word in ("Alpha ", "Beta ", "Gamma"):
            para.add_run(word)
        mgr = CommentManager(doc)
        comment_id = mgr.add_comment(para, "c", author_obj("A"), start_run=1, end_run=1)

        # Rewrite the reference run the way Word authors it: rPr + reference.
        ref = para._element.find(f".//{qn(NS_W, 'commentReference')}")
        ref_run = ref.getparent()
        rpr = etree.Element(qn(NS_W, "rPr"))
        style = etree.SubElement(rpr, qn(NS_W, "rStyle"))
        style.set(qn(NS_W, "val"), "CommentReference")
        ref_run.insert(0, rpr)

        run_count_before = len(para._element.findall(qn(NS_W, "r")))
        mgr.delete_comment(comment_id)

        runs = para._element.findall(qn(NS_W, "r"))
        assert len(runs) == run_count_before - 1
        # No ghost rPr-only run remains.
        for run in runs:
            assert [etree.QName(c).localname for c in run] != ["rPr"]

    def test_hostile_comment_id_handled_without_crash(self):
        doc = Document()
        para = doc.add_paragraph("Text")
        mgr = CommentManager(doc)
        comment_id = mgr.add_comment(para, "c", author_obj("A"))

        # Simulate a malicious/malformed document: rewrite the id everywhere
        # to contain a quote character.
        hostile = "1'1"
        for elem in mgr._comments_xml.iter(qn(NS_W, "comment")):
            if elem.get(qn(NS_W, "id")) == comment_id:
                elem.set(qn(NS_W, "id"), hostile)
        for tag in ("commentRangeStart", "commentRangeEnd", "commentReference"):
            for elem in doc.element.iter(qn(NS_W, tag)):
                if elem.get(qn(NS_W, "id")) == comment_id:
                    elem.set(qn(NS_W, "id"), hostile)
        mgr._save_comments()

        listed = list(mgr.list_comments())
        assert listed[0].comment_id == hostile
        mgr.delete_comment(hostile)
        assert list(mgr.list_comments()) == []
        assert doc.element.find(f".//{qn(NS_W, 'commentRangeStart')}") is None


class TestTableAnchors:
    def test_find_paragraph_with_comment_in_table_cell(self):
        doc = Document()
        table = doc.add_table(rows=1, cols=1)
        cell_para = table.rows[0].cells[0].paragraphs[0]
        cell_para.add_run("Cell text")
        mgr = CommentManager(doc)
        comment_id = mgr.add_comment(cell_para, "table comment", author_obj("A"))

        anchor = CommentAnchor(doc)
        found = anchor.find_paragraph_with_comment(comment_id)
        assert found is not None
        assert found._element is cell_para._element


class TestFootnoteAnchors:
    """Comments anchored in footnotes must be reachable and persist."""

    @staticmethod
    def _setup_footnote_comment(doc, mgr, comment_id):
        """Relocate a comment's anchors into a manually created footnotes part."""
        from docx.opc.packuri import PackURI
        from docx.opc.part import Part

        footnotes_xml = (
            f'<w:footnotes xmlns:w="{NS_W}"><w:footnote w:id="1"><w:p>'
            f"<w:r><w:t>note text</w:t></w:r>"
            f"</w:p></w:footnote></w:footnotes>"
        ).encode()
        part = Part(
            PackURI("/word/footnotes.xml"),
            CT_FOOTNOTES,
            footnotes_xml,
            doc.part.package,
        )
        doc.part.relate_to(part, REL_FOOTNOTES)

        # Move the anchors from the body into the footnote paragraph.
        anchor = CommentAnchor(doc)
        anchor.remove_anchors(comment_id)
        root = part_element(part)
        footnote_para = root.find(f".//{qn(NS_W, 'p')}")
        note_run = footnote_para.find(qn(NS_W, "r"))

        start = etree.Element(qn(NS_W, "commentRangeStart"))
        start.set(qn(NS_W, "id"), comment_id)
        note_run.addprevious(start)
        end = etree.Element(qn(NS_W, "commentRangeEnd"))
        end.set(qn(NS_W, "id"), comment_id)
        note_run.addnext(end)
        ref_run = etree.Element(qn(NS_W, "r"))
        ref = etree.SubElement(ref_run, qn(NS_W, "commentReference"))
        ref.set(qn(NS_W, "id"), comment_id)
        end.addnext(ref_run)
        part._blob = etree.tostring(root, xml_declaration=True, encoding="UTF-8")
        return part

    def test_reply_and_delete_for_footnote_anchored_comment(self, tmp_path):
        doc = Document()
        body_para = doc.add_paragraph("Body text")
        mgr = CommentManager(doc)
        root_id = mgr.add_comment(body_para, "on a footnote", author_obj("A"))
        self._setup_footnote_comment(doc, mgr, root_id)

        # Replying must find the footnote-hosted anchors (previously raised).
        reply_id = mgr.reply_to_comment(root_id, "reply", author_obj("B"))

        # The reply anchors must persist into the saved footnotes part.
        _, path = saved_zip_names(doc, tmp_path, "footnote.docx")
        with ZipFile(str(path)) as zf:
            footnotes = etree.fromstring(zf.read("word/footnotes.xml"))
        anchored_ids = {
            e.get(qn(NS_W, "id"))
            for e in footnotes.iter(qn(NS_W, "commentRangeStart"))
        }
        assert anchored_ids == {root_id, reply_id}

        # Deleting the thread must remove the footnote anchors, in memory and
        # in the saved file.
        mgr.delete_thread(root_id)
        _, path2 = saved_zip_names(doc, tmp_path, "footnote_deleted.docx")
        with ZipFile(str(path2)) as zf:
            footnotes = etree.fromstring(zf.read("word/footnotes.xml"))
        for tag in ("commentRangeStart", "commentRangeEnd", "commentReference"):
            assert footnotes.find(f".//{qn(NS_W, tag)}") is None


class TestLifecycleSafety:
    """Failed operations must not mutate the document."""

    @staticmethod
    def _strip_para_ids(mgr, comment_id):
        for comment_elem in mgr._comments_xml.findall(qn(NS_W, "comment")):
            if comment_elem.get(qn(NS_W, "id")) == comment_id:
                para_elem = comment_elem.find(qn(NS_W, "p"))
                para_elem.attrib.pop(qn(NS_W14, "paraId"), None)
                para_elem.attrib.pop(qn(NS_W14, "textId"), None)
                return para_elem
        raise AssertionError("comment not found")

    def test_delete_nonexistent_does_not_migrate(self):
        doc = Document()
        mgr = CommentManager(doc)
        cid = mgr.add_comment(doc.add_paragraph("Text"), "c", author_obj("A"))
        para_elem = self._strip_para_ids(mgr, cid)
        mgr._save_comments()

        with pytest.raises(ValueError, match="not found"):
            mgr.delete_comment("does-not-exist")
        # Metadata migration must not have run as a side effect.
        assert para_elem.get(qn(NS_W14, "paraId")) is None

    def test_reply_to_nonexistent_does_not_migrate(self):
        doc = Document()
        mgr = CommentManager(doc)
        cid = mgr.add_comment(doc.add_paragraph("Text"), "c", author_obj("A"))
        para_elem = self._strip_para_ids(mgr, cid)
        mgr._save_comments()

        with pytest.raises(ValueError, match="not found"):
            mgr.reply_to_comment("does-not-exist", "r", author_obj("B"))
        assert para_elem.get(qn(NS_W14, "paraId")) is None

    def test_read_only_construction_leaves_document_untouched(self, tmp_path):
        doc = Document()
        doc.add_paragraph("Pristine")
        mgr = CommentManager(doc)
        assert list(mgr.list_comments()) == []
        assert mgr.get_comment_threads() == []

        names, _ = saved_zip_names(doc, tmp_path)
        for part_name in (
            "word/comments.xml",
            "word/commentsExtended.xml",
            "word/commentsIds.xml",
            "word/commentsExtensible.xml",
        ):
            assert part_name not in names

    def test_delete_thread_unknown_id_no_mutation(self):
        doc = Document()
        para = doc.add_paragraph("text")
        mgr = CommentManager(doc)
        mgr.add_comment(para, "c", author_obj("A"))
        before = etree.tostring(doc.element.body)
        with pytest.raises(CommentNotFoundError):
            mgr.delete_thread("999999")
        assert etree.tostring(doc.element.body) == before
        assert len(list(mgr.list_comments())) == 1

    def test_move_ops_unknown_id_no_mutation(self):
        doc = Document()
        para = doc.add_paragraph("text")
        target = doc.add_paragraph("target")
        mgr = CommentManager(doc)
        mgr.add_comment(para, "c", author_obj("A"))
        before = etree.tostring(doc.element.body)
        with pytest.raises(CommentNotFoundError):
            mgr.move_comment("999999", target)
        with pytest.raises(CommentNotFoundError):
            mgr.move_thread("999999", target)
        assert etree.tostring(doc.element.body) == before

    def test_move_thread_with_explicit_indices(self):
        doc = Document()
        para = doc.add_paragraph("source")
        target = doc.add_paragraph()
        target.add_run("one ")
        target.add_run("two ")
        target.add_run("three")
        mgr = CommentManager(doc)
        cid = mgr.add_comment(para, "root", author_obj("A"))
        mgr.reply_to_comment(cid, "reply", author_obj("B"))
        mgr.move_thread(cid, target, start_run=1, end_run=1)
        assert mgr.get_anchored_text(cid) == "two "
        thread = mgr.get_thread(cid)
        assert thread.reply_count == 1

    def test_duplicate_comment_ids_warn_on_delete(self):
        doc = Document()
        mgr = CommentManager(doc)
        id1 = mgr.add_comment(doc.add_paragraph("One"), "c1", author_obj("A"))
        id2 = mgr.add_comment(doc.add_paragraph("Two"), "c2", author_obj("A"))
        for elem in mgr._comments_xml.findall(qn(NS_W, "comment")):
            if elem.get(qn(NS_W, "id")) == id2:
                elem.set(qn(NS_W, "id"), id1)
        mgr._save_comments()

        with pytest.warns(UserWarning, match="multiple comments share id"):
            mgr.delete_comment(id1)
        assert list(mgr.list_comments()) == []


class TestThreadedMove:
    def test_move_comment_on_threaded_root_raises(self):
        doc = Document()
        para1 = doc.add_paragraph("One")
        para2 = doc.add_paragraph("Two")
        mgr = CommentManager(doc)
        root_id = mgr.add_comment(para1, "Root", author_obj("A"))
        mgr.reply_to_comment(root_id, "Reply", author_obj("B"))

        with pytest.raises(ValueError, match="move_thread"):
            mgr.move_comment(root_id, para2)

        # Anchors unchanged: root still anchored at the original paragraph.
        anchor = CommentAnchor(doc)
        found = anchor.find_paragraph_with_comment(root_id)
        assert found._element is para1._element


@pytest.mark.skipif(not HAS_NATIVE_COMMENTS, reason="needs python-docx >= 1.2")
class TestPythonDocxInterop:
    def test_native_comments_api_works_after_manager_use(self):
        doc = Document()
        para = doc.add_paragraph("Text")
        mgr = CommentManager(doc)
        mgr.add_comment(para, "ours", author_obj("A"))

        # python-docx's own comments API must keep working in-session.
        native_para = doc.add_paragraph("More text")
        doc.add_comment(runs=native_para.runs, text="native", author="Nat")
        assert {c.text for c in mgr.list_comments()} == {"ours", "native"}

    def test_two_managers_on_one_document_both_persist(self, tmp_path):
        doc = Document()
        para1 = doc.add_paragraph("One")
        para2 = doc.add_paragraph("Two")

        mgr1 = CommentManager(doc)
        mgr1.add_comment(para1, "from A", author_obj("A"))
        mgr2 = CommentManager(doc)
        mgr2.add_comment(para2, "from B", author_obj("B"))

        _, path = saved_zip_names(doc, tmp_path)
        mgr3 = CommentManager(Document(str(path)))
        assert {c.text for c in mgr3.list_comments()} == {"from A", "from B"}

    def test_reply_to_native_comment(self):
        doc = Document()
        para = doc.add_paragraph("Native text")
        doc.add_comment(runs=para.runs, text="native", author="Nat")

        mgr = CommentManager(doc)
        native_id = next(iter(mgr.list_comments())).comment_id
        reply_id = mgr.reply_to_comment(native_id, "reply", author_obj("A"))
        threads = mgr.get_comment_threads()
        assert len(threads) == 1
        assert threads[0].replies[0].comment_id == reply_id


class TestOrphanHandling:
    def test_orphaned_reply_is_promoted_and_not_a_reply(self):
        doc = Document()
        para = doc.add_paragraph("Text")
        mgr = CommentManager(doc)
        root_id = mgr.add_comment(para, "Root", author_obj("A"))
        reply_id = mgr.reply_to_comment(root_id, "Reply", author_obj("B"))

        # Simulate another tool deleting the root comment element only.
        for elem in list(mgr._comments_xml):
            if (
                etree.QName(elem).localname == "comment"
                and elem.get(qn(NS_W, "id")) == root_id
            ):
                mgr._comments_xml.remove(elem)
        mgr._save_comments()

        comments = list(mgr.list_comments())
        assert len(comments) == 1
        assert comments[0].comment_id == reply_id
        assert comments[0].parent_para_id is None
        assert not comments[0].is_reply

        threads = mgr.get_comment_threads()
        assert len(threads) == 1
        assert threads[0].root.comment_id == reply_id

    def test_idless_comments_are_not_dropped_from_threads(self):
        doc = Document()
        mgr = CommentManager(doc)
        mgr._ensure_parts()
        for text in ("orphan one", "orphan two"):
            comment = etree.SubElement(mgr._comments_xml, qn(NS_W, "comment"))
            comment.set(qn(NS_W, "author"), "Legacy")
            para = etree.SubElement(comment, qn(NS_W, "p"))
            run = etree.SubElement(para, qn(NS_W, "r"))
            t = etree.SubElement(run, qn(NS_W, "t"))
            t.text = text
        mgr._save_comments()

        comments = list(mgr.list_comments())
        threads = mgr.get_comment_threads()
        assert len(comments) == 2
        assert len(threads) == 2
        assert {t.root.text for t in threads} == {"orphan one", "orphan two"}


class TestMigrationRepair:
    def test_migrate_reuses_unclaimed_para_id_and_keeps_thread(self):
        doc = Document()
        para = doc.add_paragraph("Text")
        mgr = CommentManager(doc)
        root_id = mgr.add_comment(para, "Root", author_obj("A"))
        reply_id = mgr.reply_to_comment(root_id, "Reply", author_obj("B"))

        original_reply_para_id = next(
            c.para_id for c in mgr.list_comments() if c.comment_id == reply_id
        )

        # Simulate a round trip that stripped the reply paragraph's paraId.
        for comment_elem in mgr._comments_xml.findall(qn(NS_W, "comment")):
            if comment_elem.get(qn(NS_W, "id")) == reply_id:
                para_elem = comment_elem.find(qn(NS_W, "p"))
                para_elem.attrib.pop(qn(NS_W14, "paraId"), None)
        mgr._save_comments()

        mgr.migrate_comment_metadata()

        reply = next(c for c in mgr.list_comments() if c.comment_id == reply_id)
        assert reply.para_id == original_reply_para_id
        # Thread intact: still one thread with one reply, no flattening.
        threads = mgr.get_comment_threads()
        assert len(threads) == 1
        assert threads[0].reply_count == 1
        # No dangling metadata entries left behind.
        threading = CommentsExtendedPart(doc).get_threading_info()
        assert set(threading) == {c.para_id for c in mgr.list_comments()}

    def test_migrate_cleans_orphan_metadata(self):
        doc = Document()
        para = doc.add_paragraph("Text")
        mgr = CommentManager(doc)
        mgr.add_comment(para, "Comment", author_obj("A"))

        ext_part = CommentsExtendedPart(doc)
        ext_part.add_comment_ex(para_id="DEADBEEF", parent_para_id=None, done=False)
        ids_part = CommentsIdsPart(doc)
        ids_part.add_comment_id(para_id="DEADBEEF", durable_id="FEEDFACE")

        mgr.migrate_comment_metadata()

        assert "DEADBEEF" not in CommentsExtendedPart(doc).get_threading_info()
        assert "DEADBEEF" not in CommentsIdsPart(doc).get_durable_ids()


class TestPeopleValidation:
    def test_person_personinfo_partial_presence_raises(self):
        doc = Document()
        para = doc.add_paragraph("Text")
        mgr = CommentManager(doc)
        with pytest.raises(ValueError, match="provider_id and user_id"):
            mgr.add_comment(
                para,
                "c",
                author_obj("A"),
                person=PersonInfo(author="A", provider_id="only-provider"),
            )

    def test_ensure_person_invalid_presence_creates_no_part(self, tmp_path):
        doc = Document()
        mgr = CommentManager(doc)
        with pytest.raises(ValueError, match="provider_id and user_id"):
            mgr.ensure_person("A", presence={"provider_id": "only-provider"})

        names, _ = saved_zip_names(doc, tmp_path)
        assert "word/people.xml" not in names


class TestSecureParsing:
    def test_external_entities_not_resolved(self, tmp_path):
        secret = tmp_path / "secret.txt"
        secret.write_text("TOPSECRET")
        xml = (
            f'<!DOCTYPE r [<!ENTITY x SYSTEM "file://{secret}">]><r>&x;</r>'
        ).encode()
        root = parse_xml_bytes(xml)
        text = (root.text or "") + "".join(c.tail or "" for c in root)
        assert "TOPSECRET" not in text


class TestDetachPersistence:
    """Dangling parent links must be cleared in the saved XML, not just the API view."""

    def test_delete_comment_clears_parent_link_in_saved_xml(self, tmp_path):
        doc = Document()
        para = doc.add_paragraph("Text")
        mgr = CommentManager(doc)
        root_id = mgr.add_comment(para, "Root", author_obj("A"))
        mgr.reply_to_comment(root_id, "Reply", author_obj("B"))

        mgr.delete_comment(root_id)

        _, path = saved_zip_names(doc, tmp_path, "detach.docx")
        with ZipFile(str(path)) as zf:
            ext = etree.fromstring(zf.read("word/commentsExtended.xml"))
        entries = [e for e in ext if etree.QName(e).localname == "commentEx"]
        assert len(entries) == 1
        assert entries[0].get(qn(NS_W15, "paraIdParent")) is None


class TestMigrateNoGraft:
    """Migrate must not graft orphan metadata onto an unrelated comment."""

    def test_orphan_metadata_not_grafted(self):
        doc = Document()
        doc.add_paragraph("Text")
        mgr = CommentManager(doc)
        mgr._ensure_parts()

        # Orphan metadata left behind by another tool's delete: a resolved
        # entry with a durable id, but no matching comment element.
        CommentsExtendedPart(doc).add_comment_ex(
            para_id="DEADBEEF", parent_para_id=None, done=True
        )
        CommentsIdsPart(doc).add_comment_id(para_id="DEADBEEF", durable_id="FEEDFACE")

        # An unrelated comment added without w14:paraId (e.g. by a tool that
        # does not write it).
        comment = etree.SubElement(mgr._comments_xml, qn(NS_W, "comment"))
        comment.set(qn(NS_W, "id"), "42")
        comment.set(qn(NS_W, "author"), "Other Tool")
        comment.set(qn(NS_W, "date"), "2026-08-19T10:00:00Z")
        p = etree.SubElement(comment, qn(NS_W, "p"))
        run = etree.SubElement(p, qn(NS_W, "r"))
        t = etree.SubElement(run, qn(NS_W, "t"))
        t.text = "totally unrelated new comment"
        mgr._save_comments()

        mgr.migrate_comment_metadata()

        info = next(c for c in mgr.list_comments() if c.comment_id == "42")
        assert info.para_id != "DEADBEEF"
        assert info.durable_id != "FEEDFACE"
        assert not info.is_resolved
        # The orphan entries are cleaned up rather than inherited.
        assert "DEADBEEF" not in CommentsExtendedPart(doc).get_threading_info()
        assert "DEADBEEF" not in CommentsIdsPart(doc).get_durable_ids()

    def test_stripped_para_id_with_matching_date_still_reused(self):
        # The corroborated-reuse path (dateUtc matches) keeps working; this
        # complements test_migrate_reuses_unclaimed_para_id_and_keeps_thread.
        doc = Document()
        para = doc.add_paragraph("Text")
        mgr = CommentManager(doc)
        root_id = mgr.add_comment(para, "Root", author_obj("A"))
        original = next(iter(mgr.list_comments())).para_id

        for comment_elem in mgr._comments_xml.findall(qn(NS_W, "comment")):
            if comment_elem.get(qn(NS_W, "id")) == root_id:
                comment_elem.find(qn(NS_W, "p")).attrib.pop(qn(NS_W14, "paraId"))
        mgr._save_comments()

        mgr.migrate_comment_metadata()
        assert next(iter(mgr.list_comments())).para_id == original


class TestMoveIndexStability:
    """Run indices are resolved before old anchors are removed."""

    def test_move_within_same_paragraph_targets_the_addressed_run(self):
        doc = Document()
        para = doc.add_paragraph("")
        para.add_run("one ")
        para.add_run("two")
        mgr = CommentManager(doc)
        comment_id = mgr.add_comment(para, "c", author_obj("A"), start_run=0, end_run=0)

        # Runs are now [one, ref-run, two]; index 2 addresses "two".
        mgr.move_comment(comment_id, para, start_run=2, end_run=2)

        children = list(para._element)
        names = [etree.QName(c).localname for c in children]
        start_idx = names.index("commentRangeStart")
        end_idx = names.index("commentRangeEnd")
        spanned = children[start_idx + 1 : end_idx]
        texts = [
            t.text
            for el in spanned
            for t in el.iter(qn(NS_W, "t"))
        ]
        assert texts == ["two"]

    def test_move_onto_own_reference_run_raises_before_mutation(self):
        doc = Document()
        para = doc.add_paragraph("")
        para.add_run("only run")
        mgr = CommentManager(doc)
        comment_id = mgr.add_comment(para, "c", author_obj("A"))

        # Runs are [only run, ref-run]; index 1 addresses the reference run,
        # which the move itself would remove.
        snapshot = etree.tostring(para._element)
        with pytest.raises(ValueError, match="recompute"):
            mgr.move_comment(comment_id, para, start_run=1, end_run=1)

        # The failed move mutated nothing: the paragraph XML is byte-identical.
        assert etree.tostring(para._element) == snapshot

    def test_default_in_place_move_works(self):
        doc = Document()
        para = doc.add_paragraph("Some text")
        mgr = CommentManager(doc)
        comment_id = mgr.add_comment(para, "c", author_obj("A"))

        # Whole-paragraph default move onto the paragraph the comment is
        # already anchored in must succeed (span re-resolved after removal).
        mgr.move_comment(comment_id, para)

        anchor = CommentAnchor(doc)
        found = anchor.find_paragraph_with_comment(comment_id)
        assert found is not None
        assert found._element is para._element
        starts = [
            e
            for e in para._element.iter(qn(NS_W, "commentRangeStart"))
            if e.get(qn(NS_W, "id")) == comment_id
        ]
        assert len(starts) == 1

    def test_default_in_place_move_thread_works(self):
        doc = Document()
        para = doc.add_paragraph("Some text")
        mgr = CommentManager(doc)
        root_id = mgr.add_comment(para, "Root", author_obj("A"))
        reply_id = mgr.reply_to_comment(root_id, "Reply", author_obj("B"))

        mgr.move_thread(root_id, para)

        anchor = CommentAnchor(doc)
        for cid in (root_id, reply_id):
            found = anchor.find_paragraph_with_comment(cid)
            assert found is not None
            assert found._element is para._element

    def test_move_thread_with_rootless_thread_raises_before_mutation(self):
        doc = Document()
        para = doc.add_paragraph("Some text")
        mgr = CommentManager(doc)
        root_id = mgr.add_comment(para, "Root", author_obj("A"))
        reply_id = mgr.reply_to_comment(root_id, "Reply", author_obj("B"))

        # Simulate a corrupt document: the root comment lost its w:id.
        for elem in mgr._comments_xml.findall(qn(NS_W, "comment")):
            if elem.get(qn(NS_W, "id")) == root_id:
                del elem.attrib[qn(NS_W, "id")]
        mgr._save_comments()

        with pytest.raises(ValueError, match="without a"):
            mgr.move_thread(reply_id, para)
        # Nothing was mutated: the reply is still anchored.
        anchor = CommentAnchor(doc)
        assert anchor.find_paragraph_with_comment(reply_id) is not None

    def test_detached_tree_paragraph_raises(self):
        doc = Document()
        doc.add_paragraph("Real text")
        mgr = CommentManager(doc)

        from docx.text.paragraph import Paragraph

        detached = etree.fromstring(
            f'<w:p xmlns:w="{NS_W}"><w:r><w:t>ghost</w:t></w:r></w:p>'
        )
        ghost_para = Paragraph(detached, doc)
        with pytest.raises(ValueError, match="detached"):
            mgr.add_comment(ghost_para, "c", author_obj("A"))


class TestFootnoteAnchorWrites:
    def test_move_comment_into_footnote_persists(self, tmp_path):
        doc = Document()
        body_para = doc.add_paragraph("Body text")
        other_para = doc.add_paragraph("Other text")
        mgr = CommentManager(doc)
        root_id = mgr.add_comment(body_para, "on a footnote", author_obj("A"))
        TestFootnoteAnchors._setup_footnote_comment(doc, mgr, root_id)
        standalone_id = mgr.add_comment(other_para, "standalone", author_obj("B"))

        foot_para = CommentAnchor(doc).find_paragraph_with_comment(root_id)
        assert foot_para is not None
        mgr.move_comment(standalone_id, foot_para)

        _, path = saved_zip_names(doc, tmp_path, "footnote_move.docx")
        with ZipFile(str(path)) as zf:
            footnotes = etree.fromstring(zf.read("word/footnotes.xml"))
        anchored = {
            e.get(qn(NS_W, "id"))
            for e in footnotes.iter(qn(NS_W, "commentRangeStart"))
        }
        assert standalone_id in anchored


class TestHandlerRobustness:
    def test_malformed_part_blob_raises_instead_of_dropping_writes(self):
        doc = Document()
        para = doc.add_paragraph("Text")
        mgr = CommentManager(doc)
        mgr.add_comment(para, "c", author_obj("A"))

        ext = CommentsExtendedPart(doc)
        part = ext._get_part()
        part._blob = b"this is not xml"

        with pytest.raises(ValueError, match="commentsExtended"):
            CommentsExtendedPart(doc).get_threading_info()

    def test_direct_blob_write_is_visible_to_new_handlers(self):
        doc = Document()
        para = doc.add_paragraph("Text")
        mgr = CommentManager(doc)
        mgr.add_comment(para, "c", author_obj("A"))

        ext = CommentsExtendedPart(doc)
        assert ext.get_threading_info()  # populate the part-level cache
        part = ext._get_part()
        replacement = (
            f'<w15:commentsEx xmlns:w15="{NS_W15}">'
            f'<w15:commentEx w15:paraId="AABBCCDD" w15:done="0"/>'
            f"</w15:commentsEx>"
        ).encode()
        part._blob = replacement

        info = CommentsExtendedPart(doc).get_threading_info()
        assert set(info) == {"AABBCCDD"}

    def test_external_footnotes_relationship_ignored(self):
        doc = Document()
        para = doc.add_paragraph("Text")
        doc.part.relate_to("http://example.com/notes", REL_FOOTNOTES, is_external=True)

        mgr = CommentManager(doc)
        comment_id = mgr.add_comment(para, "c", author_obj("A"))
        mgr.delete_comment(comment_id)
        assert list(mgr.list_comments()) == []


class TestRunlessParagraphIndices:
    def test_empty_paragraph_explicit_indices_raise(self):
        doc = Document()
        para = doc.add_paragraph("")  # no runs
        mgr = CommentManager(doc)
        with pytest.raises(IndexError):
            mgr.add_comment(para, "c", author_obj("A"), start_run=1)
        with pytest.raises(IndexError):
            mgr.add_comment(para, "c", author_obj("A"), end_run=0)


class TestEmptyParagraphAnchoring:
    def test_default_anchor_on_runless_paragraph(self, tmp_path):
        doc = Document()
        para = doc.add_paragraph("")
        mgr = CommentManager(doc)
        cid = mgr.add_comment(para, "c", author_obj("A"))
        locals_ = [etree.QName(c).localname for c in para._element]
        # pPr may precede (python-docx may not add one for a bare paragraph)
        anchored = [x for x in locals_ if x != "pPr"]
        assert anchored == ["commentRangeStart", "commentRangeEnd", "r"]
        path = tmp_path / "e.docx"
        doc.save(str(path))
        mgr2 = CommentManager(Document(str(path)))
        assert mgr2.get_comment(cid).text == "c"
        assert mgr2.get_comment_paragraph(cid) is not None

    def test_anchor_after_ppr_on_styled_empty_paragraph(self):
        doc = Document()
        para = doc.add_paragraph("", style="Heading 1")
        mgr = CommentManager(doc)
        mgr.add_comment(para, "c", author_obj("A"))
        locals_ = [etree.QName(c).localname for c in para._element]
        assert locals_[0] == "pPr"
        assert locals_[1] == "commentRangeStart"

    def test_move_comment_onto_empty_paragraph(self):
        doc = Document()
        para = doc.add_paragraph("source")
        empty = doc.add_paragraph("")
        mgr = CommentManager(doc)
        cid = mgr.add_comment(para, "c", author_obj("A"))
        mgr.move_comment(cid, empty)
        assert mgr.get_comment_paragraph(cid)._element is empty._element


class TestTextboxText:
    def test_textbox_text_counted_once(self):
        doc = Document()
        mgr = CommentManager(doc)
        comment_id = mgr.add_comment(doc.add_paragraph("Text"), "outer ", author_obj("A"))

        # Inject a textbox (nested paragraph) into the comment's paragraph.
        for comment_elem in mgr._comments_xml.findall(qn(NS_W, "comment")):
            if comment_elem.get(qn(NS_W, "id")) == comment_id:
                outer_p = comment_elem.find(qn(NS_W, "p"))
                holder_run = etree.SubElement(outer_p, qn(NS_W, "r"))
                pict = etree.SubElement(holder_run, qn(NS_W, "pict"))
                txbx = etree.SubElement(pict, qn(NS_W, "txbxContent"))
                inner_p = etree.SubElement(txbx, qn(NS_W, "p"))
                inner_r = etree.SubElement(inner_p, qn(NS_W, "r"))
                inner_t = etree.SubElement(inner_r, qn(NS_W, "t"))
                inner_t.text = "inner"
        mgr._save_comments()

        text = next(iter(mgr.list_comments())).text
        assert text.count("inner") == 1
        assert "outer" in text


class TestIdPoolCoversAllStories:
    def test_header_anchor_ids_excluded_from_new_ids(self, monkeypatch):
        from docx_comments import manager as manager_mod

        doc = Document()
        header_para = doc.sections[0].header.paragraphs[0]
        header_para.add_run("Header text")
        stray = etree.SubElement(
            header_para._element.getparent(), qn(NS_W, "commentRangeStart")
        )
        stray.set(qn(NS_W, "id"), "777")

        seq = iter(["777", "778"])
        monkeypatch.setattr(manager_mod, "_generate_id", lambda: next(seq))

        mgr = CommentManager(doc)
        comment_id = mgr.add_comment(doc.add_paragraph("Body"), "c", author_obj("A"))
        assert comment_id == "778"


class TestPersonValidationBeforeParts:
    def test_invalid_person_spec_creates_no_parts(self, tmp_path):
        doc = Document()
        para = doc.add_paragraph("Text")
        mgr = CommentManager(doc)
        with pytest.raises(ValueError, match="must match comment author"):
            mgr.add_comment(para, "c", author_obj("A"), person="Somebody Else")

        names, _ = saved_zip_names(doc, tmp_path)
        for part_name in (
            "word/comments.xml",
            "word/commentsExtended.xml",
            "word/commentsIds.xml",
            "word/commentsExtensible.xml",
            "word/people.xml",
        ):
            assert part_name not in names


class TestDuplicateCommentExEntries:
    def test_dangling_parent_cleared_on_every_duplicate(self, tmp_path):
        doc = Document()
        para = doc.add_paragraph("Text")
        mgr = CommentManager(doc)
        mgr.add_comment(para, "c", author_obj("A"))
        para_id = next(iter(mgr.list_comments())).para_id

        # A second commentEx for the same paraId with a dangling parent, as a
        # buggy external tool might write.
        ext = CommentsExtendedPart(doc)
        dup = etree.SubElement(ext.xml, qn(NS_W15, "commentEx"))
        dup.set(qn(NS_W15, "paraId"), para_id)
        dup.set(qn(NS_W15, "done"), "0")
        dup.set(qn(NS_W15, "paraIdParent"), "FFFFDEAD")
        ext._save()

        mgr.migrate_comment_metadata()

        _, path = saved_zip_names(doc, tmp_path, "dup_commentex.docx")
        with ZipFile(str(path)) as zf:
            saved = etree.fromstring(zf.read("word/commentsExtended.xml"))
        for entry in saved.iter(qn(NS_W15, "commentEx")):
            assert entry.get(qn(NS_W15, "paraIdParent")) is None


class TestAlternateContentText:
    def test_word_style_textbox_counted_once(self):
        NS_MC = "http://schemas.openxmlformats.org/markup-compatibility/2006"
        doc = Document()
        mgr = CommentManager(doc)
        comment_id = mgr.add_comment(doc.add_paragraph("Text"), "outer ", author_obj("A"))

        # Word serializes a textbox twice: mc:Choice (DrawingML) and
        # mc:Fallback (VML) both contain the same w:txbxContent.
        for comment_elem in mgr._comments_xml.findall(qn(NS_W, "comment")):
            if comment_elem.get(qn(NS_W, "id")) == comment_id:
                outer_p = comment_elem.find(qn(NS_W, "p"))
                holder_run = etree.SubElement(outer_p, qn(NS_W, "r"))
                alt = etree.SubElement(holder_run, qn(NS_MC, "AlternateContent"))
                for branch in ("Choice", "Fallback"):
                    wrap = etree.SubElement(alt, qn(NS_MC, branch))
                    txbx = etree.SubElement(wrap, qn(NS_W, "txbxContent"))
                    inner_p = etree.SubElement(txbx, qn(NS_W, "p"))
                    inner_r = etree.SubElement(inner_p, qn(NS_W, "r"))
                    inner_t = etree.SubElement(inner_r, qn(NS_W, "t"))
                    inner_t.text = "inner"
        mgr._save_comments()

        text = next(iter(mgr.list_comments())).text
        assert text.count("inner") == 1


class TestDuplicateCommentExResolve:
    def test_resolve_sticks_with_duplicate_commentex_entries(self):
        doc = Document()
        para = doc.add_paragraph("Text")
        mgr = CommentManager(doc)
        comment_id = mgr.add_comment(para, "c", author_obj("A"))
        para_id = next(iter(mgr.list_comments())).para_id

        # Duplicate commentEx for the same paraId, as a buggy tool might write.
        ext = CommentsExtendedPart(doc)
        dup = etree.SubElement(ext.xml, qn(NS_W15, "commentEx"))
        dup.set(qn(NS_W15, "paraId"), para_id)
        dup.set(qn(NS_W15, "done"), "0")
        ext._save()

        mgr.resolve_comment(comment_id)

        # Every duplicate carries done="1", and the read path agrees.
        entries = [
            e.get(qn(NS_W15, "done"))
            for e in CommentsExtendedPart(doc).xml
            if etree.QName(e).localname == "commentEx"
        ]
        assert entries and all(v == "1" for v in entries)
        assert next(iter(mgr.list_comments())).is_resolved


class TestBlockLevelAlternateContent:
    def test_block_level_fallback_counted_once(self):
        NS_MC = "http://schemas.openxmlformats.org/markup-compatibility/2006"
        doc = Document()
        mgr = CommentManager(doc)
        comment_id = mgr.add_comment(doc.add_paragraph("Text"), "outer", author_obj("A"))

        # Block-level mc:AlternateContent directly under w:comment, with the
        # same paragraph duplicated in Choice and Fallback.
        for comment_elem in mgr._comments_xml.findall(qn(NS_W, "comment")):
            if comment_elem.get(qn(NS_W, "id")) == comment_id:
                alt = etree.SubElement(comment_elem, qn(NS_MC, "AlternateContent"))
                for branch in ("Choice", "Fallback"):
                    wrap = etree.SubElement(alt, qn(NS_MC, branch))
                    inner_p = etree.SubElement(wrap, qn(NS_W, "p"))
                    inner_r = etree.SubElement(inner_p, qn(NS_W, "r"))
                    inner_t = etree.SubElement(inner_r, qn(NS_W, "t"))
                    inner_t.text = "inner"
        mgr._save_comments()

        text = next(iter(mgr.list_comments())).text
        assert text.count("inner") == 1


class TestMixedReferenceRuns:
    def test_move_onto_mixed_text_and_ref_run_succeeds(self):
        doc = Document()
        para = doc.add_paragraph("")
        para.add_run("alpha ")
        para.add_run("beta")
        mgr = CommentManager(doc)
        comment_id = mgr.add_comment(para, "c", author_obj("A"), start_run=0, end_run=0)

        # Merge the reference into the "beta" text run, as some foreign
        # producers do. remove_anchors keeps such mixed runs (only the ref
        # child is stripped), so a move addressing this run is safe.
        ref = para._element.find(f".//{qn(NS_W, 'commentReference')}")
        ref_run = ref.getparent()
        beta_run = [
            r
            for r in para._element.findall(qn(NS_W, "r"))
            if r.find(qn(NS_W, "t")) is not None and r.find(qn(NS_W, "t")).text == "beta"
        ][0]
        beta_run.append(ref)
        ref_run.getparent().remove(ref_run)

        runs = para._element.findall(qn(NS_W, "r"))
        beta_idx = runs.index(beta_run)
        mgr.move_comment(comment_id, para, start_run=beta_idx, end_run=beta_idx)

        # The mixed run survived with its text, and the comment is anchored
        # around it.
        children = list(para._element)
        names = [etree.QName(c).localname for c in children]
        start_idx = names.index("commentRangeStart")
        end_idx = names.index("commentRangeEnd")
        spanned_text = [
            t.text
            for el in children[start_idx + 1 : end_idx]
            for t in el.iter(qn(NS_W, "t"))
        ]
        assert spanned_text == ["beta"]


class TestDoneLexicalForms:
    """w15:done is ST_OnOff: "1"/"true"/"on" all mean resolved."""

    @pytest.mark.parametrize(
        "raw,expected",
        [("1", True), ("true", True), ("on", True), ("0", False),
         ("false", False), ("off", False), ("garbage", False), (" TRUE ", True)],
    )
    def test_st_onoff_values(self, raw, expected):
        doc = Document()
        para = doc.add_paragraph("text")
        mgr = CommentManager(doc)
        mgr.add_comment(para, "c", PersonInfo(author="A"))
        ext = CommentsExtendedPart(doc)
        for elem in ext.xml:
            if etree.QName(elem).localname == "commentEx":
                elem.set(qn(NS_W15, "done"), raw)
        ext._save()
        comment = next(iter(mgr.list_comments()))
        assert comment.is_resolved is expected


class TestMcIgnorableMigration:
    """migrate_comment_metadata declares mc:Ignorable when backfilling w14."""

    def test_migrate_adds_mc_ignorable_when_backfilling(self):
        doc = Document()
        para = doc.add_paragraph("text")
        mgr = CommentManager(doc)
        mgr.add_comment(para, "c", author_obj("A"))
        root = CommentsPart(doc).xml
        # Simulate a foreign comments.xml: prefixes declared on the root but
        # no mc:Ignorable and no w14 paragraph attributes.
        root.attrib.pop(qn(NS_MC, "Ignorable"), None)
        for p in root.iter(qn(NS_W, "p")):
            p.attrib.pop(qn(NS_W14, "paraId"), None)
            p.attrib.pop(qn(NS_W14, "textId"), None)

        mgr.migrate_comment_metadata()

        ignorable = root.get(qn(NS_MC, "Ignorable"))
        assert ignorable is not None and "w14" in ignorable.split()


class TestBlockLevelAnchors:
    def _hoist_to_body(self, doc, p_first, p_last):
        body = doc.element.body
        start = body.find(f".//{qn(NS_W, 'commentRangeStart')}")
        end = body.find(f".//{qn(NS_W, 'commentRangeEnd')}")
        p_first._element.addprevious(start)
        p_last._element.addnext(end)

    def test_reply_to_body_level_range_is_schema_valid(self, tmp_path):
        doc = Document()
        p1 = doc.add_paragraph("first")
        p2 = doc.add_paragraph("second")
        mgr = CommentManager(doc)
        cid = mgr.add_comment(p1, "root", PersonInfo(author="A"))
        self._hoist_to_body(doc, p1, p2)
        rid = mgr.reply_to_comment(cid, "reply", PersonInfo(author="B"))
        body = doc.element.body
        assert all(etree.QName(c).localname != "r" for c in body), (
            "bare w:r under w:body is schema-invalid"
        )
        path = tmp_path / "b.docx"
        doc.save(str(path))
        doc2 = Document(str(path))
        threads = CommentManager(doc2).get_comment_threads()
        assert len(threads) == 1 and threads[0].reply_count == 1
        assert threads[0].replies[0].comment_id == rid

    def test_reply_body_level_without_parent_ref_run(self):
        doc = Document()
        p1 = doc.add_paragraph("first")
        p2 = doc.add_paragraph("second")
        mgr = CommentManager(doc)
        cid = mgr.add_comment(p1, "root", PersonInfo(author="A"))
        self._hoist_to_body(doc, p1, p2)
        # Strip the parent's reference run to force the descend-into-
        # paragraph fallback.
        ref = doc.element.body.find(f".//{qn(NS_W, 'commentReference')}")
        ref_run = ref.getparent()
        ref_run.getparent().remove(ref_run)
        mgr.reply_to_comment(cid, "reply", PersonInfo(author="B"))
        body = doc.element.body
        assert all(etree.QName(c).localname != "r" for c in body)
        # Reference run landed inside the last paragraph of the range.
        assert p2._element.find(f".//{qn(NS_W, 'commentReference')}") is not None

    def test_reply_to_tr_level_range_is_schema_valid(self):
        doc = Document()
        table = doc.add_table(rows=1, cols=2)
        cell_para = table.cell(0, 0).paragraphs[0]
        cell_para.add_run("cell text")
        mgr = CommentManager(doc)
        cid = mgr.add_comment(cell_para, "root", PersonInfo(author="A"))
        tr = table._tbl.tr_lst[0]
        start = tr.find(f".//{qn(NS_W, 'commentRangeStart')}")
        end = tr.find(f".//{qn(NS_W, 'commentRangeEnd')}")
        first_tc = tr.find(qn(NS_W, "tc"))
        first_tc.addprevious(start)
        tr.append(end)
        mgr.reply_to_comment(cid, "reply", PersonInfo(author="B"))
        assert all(etree.QName(c).localname != "r" for c in tr), (
            "bare w:r under w:tr is schema-invalid"
        )

    def test_reply_to_reference_only_anchor(self, tmp_path):
        doc = Document()
        para = doc.add_paragraph("text")
        mgr = CommentManager(doc)
        cid = mgr.add_comment(para, "root", PersonInfo(author="A"))
        # Strip range markers, keep the reference (legal per ECMA-376).
        body = doc.element.body
        for tag in ("commentRangeStart", "commentRangeEnd"):
            for elem in list(body.iter(qn(NS_W, tag))):
                elem.getparent().remove(elem)
        rid = mgr.reply_to_comment(cid, "reply", PersonInfo(author="B"))
        path = tmp_path / "r.docx"
        doc.save(str(path))
        doc2 = Document(str(path))
        threads = CommentManager(doc2).get_comment_threads()
        assert len(threads) == 1 and threads[0].reply_count == 1
        # Range markers were synthesized for the reply.
        starts = [
            e for e in doc2.element.body.iter(qn(NS_W, "commentRangeStart"))
            if e.get(qn(NS_W, "id")) == rid
        ]
        assert len(starts) == 1

    def test_reply_with_no_anchors_at_all_still_raises(self):
        from docx_comments.anchors import CommentAnchor

        doc = Document()
        para = doc.add_paragraph("text")
        mgr = CommentManager(doc)
        cid = mgr.add_comment(para, "root", PersonInfo(author="A"))
        CommentAnchor(doc).remove_anchors(cid)
        with pytest.raises(ValueError, match="Could not find anchors"):
            mgr.reply_to_comment(cid, "reply", PersonInfo(author="B"))
        assert len(list(mgr.list_comments())) == 1, "no orphan reply left behind"

    def test_reply_tr_level_without_parent_ref_run_descends_into_cell(self):
        doc = Document()
        table = doc.add_table(rows=1, cols=2)
        cell_para = table.cell(0, 0).paragraphs[0]
        cell_para.add_run("cell text")
        mgr = CommentManager(doc)
        cid = mgr.add_comment(cell_para, "root", PersonInfo(author="A"))
        tr = table._tbl.tr_lst[0]
        start = tr.find(f".//{qn(NS_W, 'commentRangeStart')}")
        end = tr.find(f".//{qn(NS_W, 'commentRangeEnd')}")
        first_tc = tr.find(qn(NS_W, "tc"))
        first_tc.addprevious(start)
        tr.append(end)
        # Strip the parent's reference run to force the descend-into-
        # paragraph fallback at block level.
        ref = tr.find(f".//{qn(NS_W, 'commentReference')}")
        ref_run = ref.getparent()
        ref_run.getparent().remove(ref_run)
        mgr.reply_to_comment(cid, "reply", PersonInfo(author="B"))
        assert all(etree.QName(c).localname != "r" for c in tr), (
            "bare w:r under w:tr is schema-invalid"
        )
        # Reference run landed inside the last cell's last paragraph.
        last_cell_para = table.cell(0, 1).paragraphs[-1]
        assert (
            last_cell_para._element.find(f".//{qn(NS_W, 'commentReference')}")
            is not None
        )

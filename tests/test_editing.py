"""Tests for editing comment anchors and lifecycle operations."""

from datetime import datetime, timezone

import pytest
from docx import Document
from lxml import etree

from docx_comments import CommentManager, CommentNotFoundError, PersonInfo
from docx_comments.anchors import CommentAnchor
from docx_comments.xml_parts import (
    CommentsExtendedPart,
    CommentsExtensiblePart,
    CommentsIdsPart,
)

NS_W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
NS_W14 = "http://schemas.microsoft.com/office/word/2010/wordml"


def qn(ns: str, name: str) -> str:
    return f"{{{ns}}}{name}"


def author_obj(name: str) -> PersonInfo:
    return PersonInfo(author=name)


class TestCommentEditing:
    """Tests for comment deletion and re-anchoring behavior."""

    def test_unresolve_comment(self):
        """Resolved comments can be marked unresolved."""
        doc = Document()
        para = doc.add_paragraph("Test paragraph")
        mgr = CommentManager(doc)

        comment_id = mgr.add_comment(
            paragraph=para,
            text="Needs work",
            author=author_obj("Reviewer"),
        )

        mgr.resolve_comment(comment_id)
        mgr.unresolve_comment(comment_id)

        comments = list(mgr.list_comments())
        assert len(comments) == 1
        assert not comments[0].is_resolved

    def test_delete_comment_detaches_replies(self):
        """Deleting a root comment detaches remaining replies."""
        doc = Document()
        para = doc.add_paragraph("Test paragraph")
        mgr = CommentManager(doc)

        root_id = mgr.add_comment(
            paragraph=para,
            text="Root comment",
            author=author_obj("Author1"),
        )
        reply_id = mgr.reply_to_comment(
            parent_id=root_id,
            text="Reply comment",
            author=author_obj("Author2"),
        )

        mgr.delete_comment(root_id)

        comments = list(mgr.list_comments())
        assert len(comments) == 1
        assert comments[0].comment_id == reply_id
        assert comments[0].parent_para_id is None

    def test_delete_thread_removes_all(self):
        """Deleting a thread removes root and replies."""
        doc = Document()
        para = doc.add_paragraph("Test paragraph")
        mgr = CommentManager(doc)

        root_id = mgr.add_comment(
            paragraph=para,
            text="Root comment",
            author=author_obj("Author1"),
        )
        reply_id = mgr.reply_to_comment(
            parent_id=root_id,
            text="Reply comment",
            author=author_obj("Author2"),
        )

        mgr.delete_thread(reply_id)

        comments = list(mgr.list_comments())
        assert len(comments) == 0

        anchor = CommentAnchor(doc)
        assert anchor.find_paragraph_with_comment(root_id) is None
        assert anchor.find_paragraph_with_comment(reply_id) is None

    def test_delete_comment_cleans_orphan_metadata(self):
        """Deleting a comment cleans orphan metadata and detaches replies."""
        doc = Document()
        para = doc.add_paragraph("Test paragraph")
        mgr = CommentManager(doc)

        root_id = mgr.add_comment(
            paragraph=para,
            text="Root comment",
            author=author_obj("Author1"),
        )
        reply_id = mgr.reply_to_comment(
            parent_id=root_id,
            text="Reply comment",
            author=author_obj("Author2"),
        )

        root_para_id = next(
            comment.para_id for comment in mgr.list_comments() if comment.comment_id == root_id
        )

        ns_w = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
        ns_w14 = "http://schemas.microsoft.com/office/word/2010/wordml"
        for comment_elem in mgr._comments_xml.findall(f"{{{ns_w}}}comment"):
            if comment_elem.get(f"{{{ns_w}}}id") == root_id:
                para_elem = comment_elem.find(f"{{{ns_w}}}p")
                para_elem.attrib.pop(f"{{{ns_w14}}}paraId", None)
                para_elem.attrib.pop(f"{{{ns_w14}}}textId", None)
                break
        mgr._save_comments()

        assert root_para_id in CommentsExtendedPart(doc).get_threading_info()
        assert root_para_id in CommentsIdsPart(doc).get_durable_ids()

        mgr.delete_comment(root_id)

        comments = list(mgr.list_comments())
        assert len(comments) == 1
        assert comments[0].comment_id == reply_id
        assert comments[0].parent_para_id is None

        assert root_para_id not in CommentsExtendedPart(doc).get_threading_info()
        assert root_para_id not in CommentsIdsPart(doc).get_durable_ids()

    def test_move_comment_updates_anchor_paragraph(self):
        """Moving a comment updates its anchor location."""
        doc = Document()
        para1 = doc.add_paragraph("Paragraph one")
        para2 = doc.add_paragraph("Paragraph two")
        mgr = CommentManager(doc)

        comment_id = mgr.add_comment(
            paragraph=para1,
            text="Move me",
            author=author_obj("Author1"),
        )

        mgr.move_comment(comment_id, para2)

        anchor = CommentAnchor(doc)
        anchored_para = anchor.find_paragraph_with_comment(comment_id)
        assert anchored_para is not None
        assert anchored_para._element is para2._element

    def test_move_thread_moves_replies(self):
        """Moving a thread re-anchors replies at the new location."""
        doc = Document()
        para1 = doc.add_paragraph("Paragraph one")
        para2 = doc.add_paragraph("Paragraph two")
        mgr = CommentManager(doc)

        root_id = mgr.add_comment(
            paragraph=para1,
            text="Root comment",
            author=author_obj("Author1"),
        )
        reply_id = mgr.reply_to_comment(
            parent_id=root_id,
            text="Reply comment",
            author=author_obj("Author2"),
        )

        mgr.move_thread(root_id, para2)

        anchor = CommentAnchor(doc)
        anchored_root = anchor.find_paragraph_with_comment(root_id)
        anchored_reply = anchor.find_paragraph_with_comment(reply_id)
        assert anchored_root is not None
        assert anchored_reply is not None
        assert anchored_root._element is para2._element
        assert anchored_reply._element is para2._element


class TestEditComment:
    def _setup(self):
        doc = Document()
        para = doc.add_paragraph("text")
        mgr = CommentManager(doc)
        cid = mgr.add_comment(para, "original", PersonInfo(author="A"), initials="A")
        return doc, mgr, cid

    def test_edit_preserves_identity_and_thread(self, tmp_path):
        doc, mgr, cid = self._setup()
        rid = mgr.reply_to_comment(cid, "reply", PersonInfo(author="B"))
        mgr.resolve_comment(cid)
        before = mgr.get_comment(cid)
        mgr.edit_comment(cid, "corrected")
        after = mgr.get_comment(cid)
        assert after.text == "corrected"
        assert after.comment_id == before.comment_id
        assert after.para_id == before.para_id
        assert after.durable_id == before.durable_id
        assert after.is_resolved is True
        path = tmp_path / "e.docx"
        doc.save(str(path))
        threads = CommentManager(Document(str(path))).get_comment_threads()
        assert threads[0].root.text == "corrected"
        assert [r.comment_id for r in threads[0].replies] == [rid]

    def test_edit_changes_text_id(self):
        doc, mgr, cid = self._setup()
        para = mgr._comments_xml.find(f".//{qn(NS_W, 'p')}")
        old_text_id = para.get(qn(NS_W14, "textId"))
        mgr.edit_comment(cid, "new")
        para = mgr._comments_xml.find(f".//{qn(NS_W, 'p')}")
        assert para.get(qn(NS_W14, "textId")) != old_text_id

    def test_edit_timestamp_updates_date_and_date_utc(self):
        doc, mgr, cid = self._setup()
        stamp = datetime(2021, 5, 6, 7, 8, 9, tzinfo=timezone.utc)
        mgr.edit_comment(cid, "new", timestamp=stamp)
        info = mgr.get_comment(cid)
        assert info.timestamp == stamp

        ext = CommentsExtensiblePart(doc).get_extensible_info()
        assert ext[info.durable_id]["date_utc"] == "2021-05-06T07:08:09Z"

    def test_edit_author_and_initials(self):
        doc, mgr, cid = self._setup()
        mgr.edit_comment(cid, "new", author="B Author", initials="BA")
        info = mgr.get_comment(cid)
        assert info.author == "B Author" and info.initials == "BA"

    def test_edit_unknown_id(self):
        _, mgr, _ = self._setup()
        with pytest.raises(CommentNotFoundError):
            mgr.edit_comment("999999", "x")

    def test_edit_illegal_text_no_mutation(self):
        _, mgr, cid = self._setup()
        with pytest.raises(ValueError, match="not allowed in XML"):
            mgr.edit_comment(cid, "bad\x00text")
        assert mgr.get_comment(cid).text == "original"

    def test_edit_anchors_untouched(self):
        doc, mgr, cid = self._setup()
        before = etree.tostring(doc.element.body)
        mgr.edit_comment(cid, "new")
        assert etree.tostring(doc.element.body) == before


class TestMigrationSkip:
    """resolve/delete skip migrate_comment_metadata when metadata is complete."""

    def _spy(self, monkeypatch):
        calls = []
        original = CommentManager.migrate_comment_metadata

        def wrapper(self):
            calls.append(1)
            original(self)

        monkeypatch.setattr(CommentManager, "migrate_comment_metadata", wrapper)
        return calls

    def test_resolve_skips_migration_when_complete(self, monkeypatch):
        doc = Document()
        para = doc.add_paragraph("text")
        mgr = CommentManager(doc)
        cid = mgr.add_comment(para, "c", PersonInfo(author="A"))
        calls = self._spy(monkeypatch)
        mgr.resolve_comment(cid)
        assert calls == []
        assert mgr.get_comment(cid).is_resolved

    def test_delete_skips_migration_when_complete(self, monkeypatch):
        doc = Document()
        para = doc.add_paragraph("text")
        mgr = CommentManager(doc)
        cid = mgr.add_comment(para, "c", PersonInfo(author="A"))
        calls = self._spy(monkeypatch)
        mgr.delete_comment(cid)
        assert calls == []
        assert list(mgr.list_comments()) == []

    def test_delete_thread_skips_migration_when_complete(self, monkeypatch):
        doc = Document()
        para = doc.add_paragraph("text")
        mgr = CommentManager(doc)
        cid = mgr.add_comment(para, "c", PersonInfo(author="A"))
        mgr.reply_to_comment(cid, "r", PersonInfo(author="B"))
        calls = self._spy(monkeypatch)
        mgr.delete_thread(cid)
        assert calls == []
        assert list(mgr.list_comments()) == []

    def test_incomplete_metadata_still_migrates(self, monkeypatch):
        doc = Document()
        para = doc.add_paragraph("text")
        mgr = CommentManager(doc)
        cid = mgr.add_comment(para, "c", PersonInfo(author="A"))
        # Strip the paraId to simulate a foreign/legacy comment.
        for p in mgr._comments_xml.iter(qn(NS_W, "p")):
            p.attrib.pop(qn(NS_W14, "paraId"), None)
        calls = self._spy(monkeypatch)
        mgr.resolve_comment(cid)
        assert calls == [1]
        assert mgr.get_comment(cid).is_resolved

    def test_orphan_threading_entry_still_migrates(self, monkeypatch):
        doc = Document()
        para = doc.add_paragraph("text")
        mgr = CommentManager(doc)
        cid = mgr.add_comment(para, "c", PersonInfo(author="A"))
        CommentsExtendedPart(doc).add_comment_ex(para_id="DEADBEEF", done=False)
        calls = self._spy(monkeypatch)
        mgr.resolve_comment(cid)
        assert calls == [1]
        # The migration removes the orphan entry.
        assert "DEADBEEF" not in CommentsExtendedPart(doc).get_threading_info()

    def test_orphan_durableless_comment_id_still_migrates(self, monkeypatch):
        # A commentId element with a paraId but NO durableId is invisible to
        # get_durable_ids yet removed by the migration's orphan cleanup, so
        # the completeness check must not treat the document as complete.
        NS_W16CID = "http://schemas.microsoft.com/office/word/2016/wordml/cid"
        doc = Document()
        para = doc.add_paragraph("text")
        mgr = CommentManager(doc)
        cid = mgr.add_comment(para, "c", PersonInfo(author="A"))
        ids_part = CommentsIdsPart(doc)
        orphan = etree.SubElement(ids_part.xml, qn(NS_W16CID, "commentId"))
        orphan.set(qn(NS_W16CID, "paraId"), "DEADBEEF")
        ids_part._save()
        calls = self._spy(monkeypatch)
        mgr.resolve_comment(cid)
        assert calls == [1]
        # The migration removes the orphan element.
        assert not [
            elem
            for elem in CommentsIdsPart(doc).xml
            if elem.get(qn(NS_W16CID, "paraId")) == "DEADBEEF"
        ]

    def test_dangling_parent_still_migrates(self, monkeypatch):
        doc = Document()
        para = doc.add_paragraph("text")
        mgr = CommentManager(doc)
        cid = mgr.add_comment(para, "c", PersonInfo(author="A"))
        info = mgr.get_comment(cid)
        CommentsExtendedPart(doc).set_parent(info.para_id, "DEADBEEF")
        calls = self._spy(monkeypatch)
        mgr.resolve_comment(cid)
        assert calls == [1]

    def test_metadata_complete_false_when_parts_missing(self):
        # migrate_comment_metadata is not a no-op on a document without
        # comment parts (it creates them), so the check must return False.
        doc = Document()
        mgr = CommentManager(doc)
        assert mgr._metadata_complete() is False

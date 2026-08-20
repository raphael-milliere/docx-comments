"""Tests for people.xml identity support."""

import warnings
import zipfile

import pytest
from docx import Document
from lxml import etree

from docx_comments import CommentManager, PersonInfo

NS_W15 = "http://schemas.microsoft.com/office/word/2012/wordml"


def author_obj(name: str) -> PersonInfo:
    return PersonInfo(author=name)


def _people_entries(path):
    with zipfile.ZipFile(path) as zf:
        if "word/people.xml" not in zf.namelist():
            return []
        root = etree.fromstring(zf.read("word/people.xml"))
    return [e for e in root if etree.QName(e).localname == "person"]


class TestPeopleXml:
    """Tests for people.xml integration."""

    def test_people_xml_not_created_by_default(self, tmp_path):
        """people.xml should not be created unless requested."""
        from zipfile import ZipFile

        doc = Document()
        para = doc.add_paragraph("Test text")
        mgr = CommentManager(doc)
        mgr.add_comment(para, "Comment", author_obj("Author"))

        output_path = tmp_path / "no_people.docx"
        doc.save(str(output_path))

        with ZipFile(str(output_path), "r") as zf:
            names = zf.namelist()
            assert "word/people.xml" not in names

    def test_add_person_without_presence(self, tmp_path):
        """Creating a person entry without presenceInfo should be valid."""
        from zipfile import ZipFile

        from lxml import etree

        doc = Document()
        mgr = CommentManager(doc)
        mgr.ensure_person("Alice")

        output_path = tmp_path / "people_no_presence.docx"
        doc.save(str(output_path))

        with ZipFile(str(output_path), "r") as zf:
            xml = etree.fromstring(zf.read("word/people.xml"))

        ns_w15 = "http://schemas.microsoft.com/office/word/2012/wordml"
        people = xml.findall(f"{{{ns_w15}}}person")
        assert len(people) == 1
        assert people[0].get(f"{{{ns_w15}}}author") == "Alice"
        assert people[0].find(f"{{{ns_w15}}}presenceInfo") is None

    def test_get_person_helper(self):
        """get_person should raise when missing and return PersonInfo when present."""
        doc = Document()
        mgr = CommentManager(doc)
        with pytest.raises(KeyError):
            mgr.get_person("Alice")

        mgr.ensure_person("Alice")
        person = mgr.get_person("Alice")
        assert isinstance(person, PersonInfo)
        assert person.author == "Alice"

    def test_get_default_author_person_from_docx(self, tmp_path):
        """Resolve default author from a DOCX source."""
        source_doc = Document()
        source_doc.core_properties.author = "Alice"
        source_mgr = CommentManager(source_doc)
        source_mgr.ensure_person(
            "Alice",
            presence={"provider_id": "AD", "user_id": "user"},
        )

        source_path = tmp_path / "author_source.docx"
        source_doc.save(str(source_path))

        mgr = CommentManager(Document())
        person, initials = mgr.get_default_author_person(
            docx_path=str(source_path),
            include_presence=True,
            strict_docx=True,
        )

        assert person.author == "Alice"
        assert person.provider_id == "AD"
        assert person.user_id == "user"
        assert initials is None

    def test_get_default_author_person_strict_docx_missing(self):
        """Strict DOCX mode should raise when the file cannot be read."""
        mgr = CommentManager(Document())
        with pytest.raises(ValueError):
            mgr.get_default_author_person(
                docx_path="does-not-exist.docx",
                strict_docx=True,
            )

    def test_get_default_author_person_warns_on_multiple_people(self, tmp_path, monkeypatch):
        """Multiple people entries should warn and fall back."""
        import docx_comments.system_author as system_author

        source_doc = Document()
        source_mgr = CommentManager(source_doc)
        source_mgr.ensure_person("Alice")
        source_mgr.ensure_person("Bob")

        source_path = tmp_path / "author_source_multi.docx"
        source_doc.save(str(source_path))

        target_doc = Document()
        target_doc.core_properties.author = "Fallback"
        target_mgr = CommentManager(target_doc)

        monkeypatch.setattr(system_author, "_system_office_user_info", lambda: (None, None))

        with warnings.catch_warnings(record=True) as caught:
            warnings.simplefilter("always")
            person, _ = target_mgr.get_default_author_person(docx_path=str(source_path))

        assert any("people entries" in str(w.message) for w in caught)
        assert person.author == "Fallback"

    def test_existing_people_xml_preserved_on_comment(self, tmp_path):
        """Adding a comment should not remove existing people.xml data."""
        from zipfile import ZipFile

        from lxml import etree

        doc = Document()
        para = doc.add_paragraph("Test text")
        mgr = CommentManager(doc)
        mgr.ensure_person("Alice")

        mgr.add_comment(para, "Comment", author_obj("Alice"))

        output_path = tmp_path / "people_preserved.docx"
        doc.save(str(output_path))

        with ZipFile(str(output_path), "r") as zf:
            xml = etree.fromstring(zf.read("word/people.xml"))

        ns_w15 = "http://schemas.microsoft.com/office/word/2012/wordml"
        people = xml.findall(f"{{{ns_w15}}}person")
        authors = {p.get(f"{{{ns_w15}}}author") for p in people}
        assert "Alice" in authors

    def test_merge_people_from_document(self):
        """Merging people should add missing authors without overwriting."""
        source_doc = Document()
        source_mgr = CommentManager(source_doc)
        source_mgr.ensure_person(
            "Alice",
            presence={"provider_id": "provider", "user_id": "user"},
        )
        source_mgr.ensure_person("Bob")

        target_doc = Document()
        target_mgr = CommentManager(target_doc)
        target_mgr.ensure_person("Bob")

        added = target_mgr.merge_people_from(source_doc)
        assert [person.author for person in added] == ["Alice"]

        people = target_mgr.get_people()
        authors = {person.author for person in people}
        assert authors == {"Alice", "Bob"}

        alice = next(person for person in people if person.author == "Alice")
        assert alice.provider_id is None
        assert alice.user_id is None

        target_doc_with_presence = Document()
        target_mgr_with_presence = CommentManager(target_doc_with_presence)
        target_mgr_with_presence.merge_people_from(source_doc, include_presence=True)
        people_with_presence = target_mgr_with_presence.get_people()
        alice_with_presence = next(
            person for person in people_with_presence if person.author == "Alice"
        )
        assert alice_with_presence.provider_id == "provider"
        assert alice_with_presence.user_id == "user"

    def test_add_comment_with_author_personinfo(self):
        """Author can be provided as PersonInfo."""
        doc = Document()
        para = doc.add_paragraph("Test text")
        mgr = CommentManager(doc)
        mgr.ensure_person("Alice")

        person = mgr.get_people()[0]
        mgr.add_comment(para, "Comment", author=person)

        comments = list(mgr.list_comments())
        assert comments[0].author == "Alice"

    def test_add_comment_with_author_presence_personinfo(self, tmp_path):
        """Author PersonInfo with presence should create people.xml entry."""
        from zipfile import ZipFile

        from lxml import etree

        doc = Document()
        para = doc.add_paragraph("Test text")
        mgr = CommentManager(doc)

        person = PersonInfo(author="Alice", provider_id="provider", user_id="user")
        mgr.add_comment(
            para,
            "Comment",
            author=person,
            person=person,
        )

        output_path = tmp_path / "author_personinfo_presence.docx"
        doc.save(str(output_path))

        with ZipFile(str(output_path), "r") as zf:
            xml = etree.fromstring(zf.read("word/people.xml"))

        ns_w15 = "http://schemas.microsoft.com/office/word/2012/wordml"
        person_elem = xml.find(f"{{{ns_w15}}}person")
        assert person_elem is not None
        assert person_elem.get(f"{{{ns_w15}}}author") == "Alice"
        presence = person_elem.find(f"{{{ns_w15}}}presenceInfo")
        assert presence is not None
        assert presence.get(f"{{{ns_w15}}}providerId") == "provider"
        assert presence.get(f"{{{ns_w15}}}userId") == "user"


class TestPersonSpecForms:
    """The README-documented person= spec forms all reach people.xml."""

    def _mgr(self):
        doc = Document()
        para = doc.add_paragraph("text")
        return doc, para, CommentManager(doc)

    def test_person_true_creates_entry(self, tmp_path):
        doc, para, mgr = self._mgr()
        mgr.add_comment(para, "c", PersonInfo(author="A"), person=True)
        path = tmp_path / "p.docx"
        doc.save(str(path))
        entries = _people_entries(path)
        assert len(entries) == 1
        assert entries[0].get(f"{{{NS_W15}}}author") == "A"

    @pytest.mark.parametrize(
        "spec",
        [
            {"provider_id": "AD", "user_id": "S::u"},
            {"providerId": "AD", "userId": "S::u"},
            {"presence": {"provider_id": "AD", "user_id": "S::u"}},
        ],
    )
    def test_person_dict_forms(self, tmp_path, spec):
        doc, para, mgr = self._mgr()
        mgr.add_comment(para, "c", PersonInfo(author="A"), person=spec)
        path = tmp_path / "p.docx"
        doc.save(str(path))
        entries = _people_entries(path)
        assert len(entries) == 1
        presence = entries[0].find(f"{{{NS_W15}}}presenceInfo")
        assert presence is not None
        assert presence.get(f"{{{NS_W15}}}providerId") == "AD"
        assert presence.get(f"{{{NS_W15}}}userId") == "S::u"

    def test_person_dict_with_author_key(self, tmp_path):
        doc, para, mgr = self._mgr()
        mgr.add_comment(para, "c", PersonInfo(author="A"), person={"author": "A"})
        path = tmp_path / "p.docx"
        doc.save(str(path))
        entries = _people_entries(path)
        assert len(entries) == 1
        assert entries[0].get(f"{{{NS_W15}}}author") == "A"

    def test_person_dict_partial_presence_raises(self):
        _, para, mgr = self._mgr()
        with pytest.raises(ValueError, match="provider_id and user_id"):
            mgr.add_comment(para, "c", PersonInfo(author="A"), person={"provider_id": "AD"})

    def test_person_invalid_type(self):
        _, para, mgr = self._mgr()
        with pytest.raises(TypeError, match="person must be"):
            mgr.add_comment(para, "c", PersonInfo(author="A"), person=42)

    def test_reply_person_true(self, tmp_path):
        doc, para, mgr = self._mgr()
        cid = mgr.add_comment(para, "c", PersonInfo(author="A"))
        mgr.reply_to_comment(cid, "r", PersonInfo(author="B"), person=True)
        path = tmp_path / "p.docx"
        doc.save(str(path))
        authors = {e.get(f"{{{NS_W15}}}author") for e in _people_entries(path)}
        assert authors == {"B"}

    def test_presence_author_auto_links(self, tmp_path):
        doc, para, mgr = self._mgr()
        mgr.add_comment(
            para,
            "c",
            PersonInfo(author="A", provider_id="AD", user_id="S::u"),
        )
        path = tmp_path / "p.docx"
        doc.save(str(path))
        assert len(_people_entries(path)) == 1

    def test_person_false_suppresses_auto_link(self, tmp_path):
        doc, para, mgr = self._mgr()
        mgr.add_comment(
            para,
            "c",
            PersonInfo(author="A", provider_id="AD", user_id="S::u"),
            person=False,
        )
        path = tmp_path / "p.docx"
        doc.save(str(path))
        assert _people_entries(path) == []


class TestPersonXmlLegality:
    def test_illegal_presence_raises_before_any_mutation(self, tmp_path):
        doc = Document()
        para = doc.add_paragraph("text")
        mgr = CommentManager(doc)
        with pytest.raises(ValueError, match="not allowed in XML"):
            mgr.add_comment(
                para,
                "c",
                PersonInfo(author="A"),
                person={"provider_id": "AD\x00", "user_id": "u"},
            )
        path = tmp_path / "clean.docx"
        doc.save(str(path))
        with zipfile.ZipFile(path) as zf:
            names = set(zf.namelist())
        assert "word/comments.xml" not in names
        assert "word/people.xml" not in names

    def test_illegal_author_in_ensure_person_no_ghost_entry(self, tmp_path):
        doc = Document()
        mgr = CommentManager(doc)
        with pytest.raises(ValueError, match="not allowed in XML"):
            mgr.ensure_person("bad\x00author")
        assert mgr.get_people() == []
        # A later valid write must not resurrect a half-built entry.
        mgr.ensure_person("Good Author")
        path = tmp_path / "p.docx"
        doc.save(str(path))
        doc2 = Document(str(path))
        people = CommentManager(doc2).get_people()
        assert [p.author for p in people] == ["Good Author"]

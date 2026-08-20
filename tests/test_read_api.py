"""Read-side API: get_comment, get_thread, anchors introspection."""

import pytest
from docx import Document
from lxml import etree

from docx_comments import CommentManager, CommentNotFoundError, PersonInfo

NS_W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def qn(ns, name):
    return f"{{{ns}}}{name}"


def _mgr_with_comment(text="hello world"):
    doc = Document()
    para = doc.add_paragraph(text)
    mgr = CommentManager(doc)
    cid = mgr.add_comment(para, "note", PersonInfo(author="A"))
    return doc, para, mgr, cid


class TestGetCommentAndThread:
    def test_get_comment(self):
        _, _, mgr, cid = _mgr_with_comment()
        info = mgr.get_comment(cid)
        assert info.comment_id == cid and info.text == "note"

    def test_get_comment_int_id(self):
        _, _, mgr, cid = _mgr_with_comment()
        assert mgr.get_comment(int(cid)).comment_id == cid

    def test_get_comment_not_found(self):
        _, _, mgr, _ = _mgr_with_comment()
        with pytest.raises(CommentNotFoundError):
            mgr.get_comment("999999")

    def test_get_thread_by_reply_id(self):
        _, _, mgr, cid = _mgr_with_comment()
        rid = mgr.reply_to_comment(cid, "r", PersonInfo(author="B"))
        thread = mgr.get_thread(rid)
        assert thread.root.comment_id == cid
        assert [c.comment_id for c in thread.replies] == [rid]

    def test_get_thread_not_found(self):
        _, _, mgr, _ = _mgr_with_comment()
        with pytest.raises(CommentNotFoundError):
            mgr.get_thread("999999")


class TestAnchorIntrospection:
    def test_get_comment_paragraph(self):
        _, para, mgr, cid = _mgr_with_comment()
        found = mgr.get_comment_paragraph(cid)
        assert found is not None and found._element is para._element

    def test_get_comment_paragraph_unknown_id(self):
        _, _, mgr, _ = _mgr_with_comment()
        with pytest.raises(CommentNotFoundError):
            mgr.get_comment_paragraph("999999")

    def test_anchored_text_whole_paragraph(self):
        _, _, mgr, cid = _mgr_with_comment("hello world")
        assert mgr.get_anchored_text(cid) == "hello world"

    def test_anchored_text_run_range(self):
        doc = Document()
        para = doc.add_paragraph()
        para.add_run("one ")
        para.add_run("two ")
        para.add_run("three")
        mgr = CommentManager(doc)
        cid = mgr.add_comment(para, "c", PersonInfo(author="A"), start_run=1, end_run=1)
        assert mgr.get_anchored_text(cid) == "two "

    def test_anchored_text_includes_container_content(self):
        doc = Document()
        para = doc.add_paragraph()
        para.add_run("See ")
        hyperlink = etree.SubElement(para._element, qn(NS_W, "hyperlink"))
        run = etree.SubElement(hyperlink, qn(NS_W, "r"))
        t = etree.SubElement(run, qn(NS_W, "t"))
        t.text = "the link"
        mgr = CommentManager(doc)
        cid = mgr.add_comment(para, "c", PersonInfo(author="A"))
        assert mgr.get_anchored_text(cid) == "See the link"

    def test_anchored_text_multi_paragraph_block_range(self):
        doc = Document()
        p1 = doc.add_paragraph("first")
        p2 = doc.add_paragraph("second")
        mgr = CommentManager(doc)
        cid = mgr.add_comment(p1, "c", PersonInfo(author="A"))
        body = doc.element.body
        start = body.find(f".//{qn(NS_W, 'commentRangeStart')}")
        end = body.find(f".//{qn(NS_W, 'commentRangeEnd')}")
        p1._element.addprevious(start)
        p2._element.addnext(end)
        assert mgr.get_anchored_text(cid) == "first\nsecond"

    def test_anchored_text_empty_paragraph(self):
        doc = Document()
        para = doc.add_paragraph("")
        mgr = CommentManager(doc)
        cid = mgr.add_comment(para, "c", PersonInfo(author="A"))
        assert mgr.get_anchored_text(cid) == ""

    def test_anchored_text_none_without_range(self):
        doc, _, mgr, cid = _mgr_with_comment()
        body = doc.element.body
        for tag in ("commentRangeStart", "commentRangeEnd"):
            for elem in list(body.iter(qn(NS_W, tag))):
                elem.getparent().remove(elem)
        assert mgr.get_anchored_text(cid) is None

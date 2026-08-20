"""Rich comment content: formatted runs and multiple paragraphs."""

import pytest
from docx import Document

from docx_comments import CommentManager, PersonInfo
from docx_comments.xml_parts import CommentsExtendedPart

NS_W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
NS_W14 = "http://schemas.microsoft.com/office/word/2010/wordml"
NS_W15 = "http://schemas.microsoft.com/office/word/2012/wordml"


def qn(ns, name):
    return f"{{{ns}}}{name}"


def _setup():
    doc = Document()
    para = doc.add_paragraph("text")
    return doc, para, CommentManager(doc)


class TestFormattedRuns:
    def test_bold_italic_underline(self):
        doc, para, mgr = _setup()
        mgr.add_comment(
            para,
            [
                [
                    ("b", {"bold": True}),
                    ("i", {"italic": True}),
                    ("u", {"underline": True}),
                    " plain",
                ]
            ],
            PersonInfo(author="A"),
        )
        comment = mgr._comments_xml.find(qn(NS_W, "comment"))
        runs = [r for r in comment.iter(qn(NS_W, "r")) if r.find(qn(NS_W, "annotationRef")) is None]
        assert runs[0].find(f"{qn(NS_W, 'rPr')}/{qn(NS_W, 'b')}") is not None
        assert runs[1].find(f"{qn(NS_W, 'rPr')}/{qn(NS_W, 'i')}") is not None
        u = runs[2].find(f"{qn(NS_W, 'rPr')}/{qn(NS_W, 'u')}")
        assert u is not None and u.get(qn(NS_W, "val")) == "single"
        assert runs[3].find(qn(NS_W, "rPr")) is None

    def test_unknown_format_key_raises(self):
        _, para, mgr = _setup()
        with pytest.raises(ValueError, match="unsupported run formatting"):
            mgr.add_comment(para, [[("x", {"blink": True})]], PersonInfo(author="A"))

    def test_plain_str_unchanged(self):
        doc, para, mgr = _setup()
        cid = mgr.add_comment(para, "line1\nline2", PersonInfo(author="A"))
        assert mgr.get_comment(cid).text == "line1\nline2"
        comment = mgr._comments_xml.find(qn(NS_W, "comment"))
        assert len(comment.findall(qn(NS_W, "p"))) == 1


class TestMultiParagraph:
    def test_two_paragraphs_keyed_to_last(self, tmp_path):
        doc, para, mgr = _setup()
        cid = mgr.add_comment(para, ["para one", "para two"], PersonInfo(author="A"))
        comment = mgr._comments_xml.find(qn(NS_W, "comment"))
        paras = comment.findall(qn(NS_W, "p"))
        assert len(paras) == 2
        pids = [p.get(qn(NS_W14, "paraId")) for p in paras]
        assert all(pids) and pids[0] != pids[1]
        threading = CommentsExtendedPart(doc).get_threading_info()
        assert pids[1] in threading and pids[0] not in threading
        assert mgr.get_comment(cid).text == "para one\npara two"
        path = tmp_path / "m.docx"
        doc.save(str(path))
        reread = CommentManager(Document(str(path))).get_comment(cid)
        assert reread.text == "para one\npara two"
        assert reread.para_id == pids[1]

    def test_annotation_ref_on_first_paragraph_only(self):
        doc, para, mgr = _setup()
        mgr.add_comment(para, ["one", "two"], PersonInfo(author="A"))
        comment = mgr._comments_xml.find(qn(NS_W, "comment"))
        paras = comment.findall(qn(NS_W, "p"))
        assert paras[0].find(f".//{qn(NS_W, 'annotationRef')}") is not None
        assert paras[1].find(f".//{qn(NS_W, 'annotationRef')}") is None

    def test_reply_and_edit_accept_rich_content(self):
        doc, para, mgr = _setup()
        cid = mgr.add_comment(para, "root", PersonInfo(author="A"))
        rid = mgr.reply_to_comment(cid, [[("hot", {"bold": True})]], PersonInfo(author="B"))
        assert mgr.get_comment(rid).text == "hot"
        mgr.edit_comment(cid, ["new one", "new two"])
        assert mgr.get_comment(cid).text == "new one\nnew two"

    def test_empty_content_raises(self):
        _, para, mgr = _setup()
        with pytest.raises(ValueError, match="at least one paragraph"):
            mgr.add_comment(para, [], PersonInfo(author="A"))

    def test_illegal_char_in_run_no_mutation(self, tmp_path):
        doc, para, mgr = _setup()
        with pytest.raises(ValueError, match="not allowed in XML"):
            mgr.add_comment(para, [["ok", ("bad\x00", {})]], PersonInfo(author="A"))
        path = tmp_path / "c.docx"
        doc.save(str(path))
        import zipfile

        with zipfile.ZipFile(path) as zf:
            assert "word/comments.xml" not in zf.namelist()

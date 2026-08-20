"""Direct tests for the system/default author resolution helpers."""

import plistlib
import sys
from zipfile import ZipFile

import pytest
from docx import Document

from docx_comments import CommentManager, system_author
from docx_comments.system_author import (
    _default_person_from_system,
    _macos_office_user_info,
    _person_from_docx,
    _system_office_user_info,
    _windows_office_user_info,
)


def _write_plist(tmp_path, data):
    office_dir = tmp_path / "Library/Group Containers/UBF8T346G9.Office"
    office_dir.mkdir(parents=True)
    with (office_dir / "MeContact.plist").open("wb") as handle:
        plistlib.dump(data, handle)


class TestMacosPlist:
    def test_reads_name_and_initials(self, tmp_path, monkeypatch):
        _write_plist(tmp_path, {"Name": "Jane Doe", "Initials": "JD"})
        monkeypatch.setattr(system_author.Path, "home", classmethod(lambda cls: tmp_path))
        assert _macos_office_user_info() == ("Jane Doe", "JD")

    def test_wrong_types_become_none(self, tmp_path, monkeypatch):
        _write_plist(tmp_path, {"Name": 42, "Initials": ["J"]})
        monkeypatch.setattr(system_author.Path, "home", classmethod(lambda cls: tmp_path))
        assert _macos_office_user_info() == (None, None)

    def test_missing_file(self, tmp_path, monkeypatch):
        monkeypatch.setattr(system_author.Path, "home", classmethod(lambda cls: tmp_path))
        assert _macos_office_user_info() == (None, None)

    def test_corrupt_plist_returns_none(self, tmp_path, monkeypatch):
        office_dir = tmp_path / "Library/Group Containers/UBF8T346G9.Office"
        office_dir.mkdir(parents=True)
        (office_dir / "MeContact.plist").write_bytes(b"not a plist")
        monkeypatch.setattr(system_author.Path, "home", classmethod(lambda cls: tmp_path))
        assert _macos_office_user_info() == (None, None)

    def test_non_dict_plist_returns_none(self, tmp_path, monkeypatch):
        _write_plist(tmp_path, ["not", "a", "dict"])
        monkeypatch.setattr(system_author.Path, "home", classmethod(lambda cls: tmp_path))
        assert _macos_office_user_info() == (None, None)


class TestPlatformDispatch:
    """Exercise the real _system_office_user_info (the autouse fixture stubs
    the module attribute, but the direct import above keeps the original)."""

    def test_unknown_platform_returns_none(self, monkeypatch):
        monkeypatch.setattr(system_author.sys, "platform", "linux")
        assert _system_office_user_info() == (None, None)

    def test_darwin_dispatches_to_macos_reader(self, tmp_path, monkeypatch):
        monkeypatch.setattr(system_author.sys, "platform", "darwin")
        monkeypatch.setattr(system_author.Path, "home", classmethod(lambda cls: tmp_path))
        assert _system_office_user_info() == (None, None)

    @pytest.mark.skipif(sys.platform == "win32", reason="reads the real registry on Windows")
    def test_windows_reader_degrades_without_winreg(self, monkeypatch):
        monkeypatch.setattr(system_author.sys, "platform", "win32")
        assert _system_office_user_info() == (None, None)

    @pytest.mark.skipif(sys.platform == "win32", reason="exercises the non-Windows early return")
    def test_windows_reader_early_return_off_windows(self):
        assert _windows_office_user_info() == (None, None)


def _author_docx(tmp_path, author="Env Author"):
    doc = Document()
    mgr = CommentManager(doc)
    mgr.ensure_person(author)
    path = tmp_path / "author-source.docx"
    doc.save(str(path))
    return path


def _replace_people_xml(src_path, dst_path, people_bytes):
    with ZipFile(src_path) as src, ZipFile(dst_path, "w") as dst:
        for item in src.infolist():
            data = src.read(item.filename)
            if item.filename == "word/people.xml":
                data = people_bytes
            dst.writestr(item, data)


class TestPersonFromDocx:
    def test_empty_path_returns_none(self):
        assert _person_from_docx("") == (None, None)

    def test_non_zip_file_returns_none(self, tmp_path):
        path = tmp_path / "not-a-zip.docx"
        path.write_bytes(b"garbage")
        assert _person_from_docx(str(path)) == (None, None)

    def test_invalid_people_xml_warns_and_falls_back(self, tmp_path):
        broken = tmp_path / "broken.docx"
        _replace_people_xml(_author_docx(tmp_path), broken, b"<not-xml")
        with pytest.warns(UserWarning, match="invalid people.xml"):
            person, _ = _default_person_from_system(docx_path=str(broken))
        assert person is None

    def test_person_without_author_name_warns_and_falls_back(self, tmp_path):
        nameless = tmp_path / "nameless.docx"
        people = (
            b"<w15:people xmlns:w15="
            b'"http://schemas.microsoft.com/office/word/2012/wordml">'
            b"<w15:person/></w15:people>"
        )
        _replace_people_xml(_author_docx(tmp_path), nameless, people)
        with pytest.warns(UserWarning, match="no author name"):
            person, _ = _default_person_from_system(docx_path=str(nameless))
        assert person is None

    def test_include_presence_without_presence_info(self, tmp_path):
        person, _ = _person_from_docx(str(_author_docx(tmp_path)), include_presence=True)
        assert person is not None and person.author == "Env Author"
        assert person.provider_id is None and person.user_id is None


class TestEnvVarSource:
    def test_env_var_used(self, tmp_path, monkeypatch):
        path = _author_docx(tmp_path)
        monkeypatch.setenv("DOCX_COMMENTS_AUTHOR_DOCX", str(path))
        person, _ = _default_person_from_system()
        assert person is not None and person.author == "Env Author"

    def test_docx_without_people_warns_and_falls_back(self, tmp_path, monkeypatch):
        doc = Document()
        path = tmp_path / "plain.docx"
        doc.save(str(path))
        with pytest.warns(UserWarning, match="no people.xml"):
            person, _ = _default_person_from_system(docx_path=str(path))
        assert person is None


class TestSystemInfoFallback:
    def test_system_office_info_used(self, monkeypatch):
        monkeypatch.setattr(system_author, "_system_office_user_info", lambda: ("Sys User", "SU"))
        person, initials = _default_person_from_system()
        assert person is not None and person.author == "Sys User"
        assert initials == "SU"


class TestManagerDefaultAuthor:
    def test_no_default_author_raises(self):
        doc = Document()
        doc.core_properties.author = ""
        doc.core_properties.last_modified_by = ""
        mgr = CommentManager(doc)
        with pytest.raises(ValueError, match="no default author"):
            mgr.get_default_author_person()

    def test_core_properties_fallback(self):
        doc = Document()
        doc.core_properties.author = "Core Author"
        mgr = CommentManager(doc)
        person, initials = mgr.get_default_author_person()
        assert person.author == "Core Author" and initials is None

    def test_last_modified_by_fallback(self):
        doc = Document()
        doc.core_properties.author = ""
        doc.core_properties.last_modified_by = "Modifier"
        mgr = CommentManager(doc)
        person, _ = mgr.get_default_author_person()
        assert person.author == "Modifier"

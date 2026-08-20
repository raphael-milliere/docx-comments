"""API polish: exception types, int-id coercion, honest type hints."""

import pytest
from docx import Document

from docx_comments import (
    CommentManager,
    CommentNotFoundError,
    PersonInfo,
    PersonNotFoundError,
)


def _doc_with_comment():
    doc = Document()
    para = doc.add_paragraph("some text")
    mgr = CommentManager(doc)
    cid = mgr.add_comment(para, "note", PersonInfo(author="A"), initials="A")
    return doc, mgr, cid


class TestExceptionTypes:
    def test_not_found_is_comment_not_found_error(self):
        _, mgr, _ = _doc_with_comment()
        with pytest.raises(CommentNotFoundError):
            mgr.resolve_comment("999999")

    def test_not_found_still_caught_as_value_error(self):
        _, mgr, _ = _doc_with_comment()
        with pytest.raises(ValueError):
            mgr.delete_comment("999999")

    def test_not_found_caught_as_lookup_error(self):
        _, mgr, _ = _doc_with_comment()
        with pytest.raises(LookupError):
            mgr.delete_thread("999999")

    @pytest.mark.parametrize(
        "op",
        ["delete_comment", "delete_thread", "resolve_comment", "unresolve_comment"],
    )
    def test_all_lookup_ops_raise_comment_not_found(self, op):
        _, mgr, _ = _doc_with_comment()
        with pytest.raises(CommentNotFoundError):
            getattr(mgr, op)("424242")

    def test_move_ops_raise_comment_not_found(self):
        doc, mgr, _ = _doc_with_comment()
        target = doc.add_paragraph("target")
        with pytest.raises(CommentNotFoundError):
            mgr.move_comment("424242", target)
        with pytest.raises(CommentNotFoundError):
            mgr.move_thread("424242", target)

    def test_reply_to_missing_parent_raises_comment_not_found(self):
        _, mgr, _ = _doc_with_comment()
        with pytest.raises(CommentNotFoundError):
            mgr.reply_to_comment("424242", "r", PersonInfo(author="B"))

    def test_person_not_found_error(self):
        _, mgr, _ = _doc_with_comment()
        with pytest.raises(PersonNotFoundError):
            mgr.get_person("Nobody")
        with pytest.raises(KeyError):
            mgr.get_person("Nobody")


class TestIntIdCoercion:
    def test_int_id_accepted_everywhere(self):
        doc, mgr, cid = _doc_with_comment()
        int_id = int(cid)
        rid = mgr.reply_to_comment(int_id, "reply", PersonInfo(author="B"))
        mgr.resolve_comment(int(rid))
        mgr.unresolve_comment(int_id)
        mgr.delete_thread(int_id)
        assert list(mgr.list_comments()) == []

    def test_bool_id_raises_type_error(self):
        _, mgr, _ = _doc_with_comment()
        with pytest.raises(TypeError):
            mgr.resolve_comment(True)

    def test_other_types_raise_type_error(self):
        _, mgr, _ = _doc_with_comment()
        with pytest.raises(TypeError):
            mgr.delete_comment(3.14)


class TestHonestHints:
    def test_comment_id_optional_annotation(self):
        from typing import Optional, get_type_hints

        from docx_comments.models import CommentInfo

        hints = get_type_hints(CommentInfo)
        assert hints["comment_id"] == Optional[str]


class TestStrAuthor:
    def test_plain_str_author(self):
        doc = Document()
        para = doc.add_paragraph("text")
        mgr = CommentManager(doc)
        cid = mgr.add_comment(para, "c", "Reviewer Name", initials="RN")
        info = mgr.get_comment(cid)
        assert info.author == "Reviewer Name" and info.initials == "RN"

    def test_str_author_in_reply(self):
        doc = Document()
        para = doc.add_paragraph("text")
        mgr = CommentManager(doc)
        cid = mgr.add_comment(para, "c", "A")
        rid = mgr.reply_to_comment(cid, "r", "B")
        assert mgr.get_comment(rid).author == "B"

    def test_str_author_on_text(self):
        doc = Document()
        para = doc.add_paragraph("find this word here")
        mgr = CommentManager(doc)
        cid = mgr.add_comment_on_text(para, "word", "c", "C")
        assert mgr.get_comment(cid).author == "C"

    def test_invalid_author_type(self):
        doc = Document()
        para = doc.add_paragraph("text")
        mgr = CommentManager(doc)
        with pytest.raises(TypeError):
            mgr.add_comment(para, "c", 42)

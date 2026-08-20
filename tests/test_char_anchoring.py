"""Character-offset and substring anchoring."""

import re

import pytest
from docx import Document
from lxml import etree

from docx_comments import CommentManager, PersonInfo
from docx_comments.anchors import CommentAnchor

NS_W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def qn(ns, name):
    return f"{{{ns}}}{name}"


def _setup(text="The quick brown fox"):
    doc = Document()
    para = doc.add_paragraph(text)
    return doc, para, CommentManager(doc)


class TestCharSpan:
    def test_mid_run_anchor(self):
        doc, para, mgr = _setup()
        cid = mgr.add_comment(
            para, "c", PersonInfo(author="A"), start_char=4, end_char=9
        )
        assert mgr.get_anchored_text(cid) == "quick"
        assert para.text == "The quick brown fox"

    def test_text_survives_save_reload(self, tmp_path):
        doc, para, mgr = _setup()
        mgr.add_comment(para, "c", PersonInfo(author="A"), start_char=4, end_char=9)
        path = tmp_path / "s.docx"
        doc.save(str(path))
        doc2 = Document(str(path))
        assert doc2.paragraphs[0].text == "The quick brown fox"

    def test_boundary_offsets_do_not_split(self):
        doc = Document()
        para = doc.add_paragraph()
        para.add_run("one ")
        para.add_run("two")
        mgr = CommentManager(doc)
        cid = mgr.add_comment(
            para, "c", PersonInfo(author="A"), start_char=4, end_char=7
        )
        assert mgr.get_anchored_text(cid) == "two"
        assert len(para._element.findall(qn(NS_W, "r"))) == 3  # 2 text + 1 ref

    def test_split_preserves_formatting(self):
        doc = Document()
        para = doc.add_paragraph()
        run = para.add_run("boldtext")
        run.bold = True
        mgr = CommentManager(doc)
        cid = mgr.add_comment(
            para, "c", PersonInfo(author="A"), start_char=4, end_char=8
        )
        assert mgr.get_anchored_text(cid) == "text"
        for r in para._element.findall(qn(NS_W, "r")):
            t = r.find(qn(NS_W, "t"))
            if t is not None and t.text:
                assert r.find(f"{qn(NS_W, 'rPr')}/{qn(NS_W, 'b')}") is not None

    def test_whitespace_edges_preserved(self):
        doc, para, mgr = _setup("keep  spaces  here")
        mgr.add_comment(para, "c", PersonInfo(author="A"), start_char=6, end_char=12)
        assert para.text == "keep  spaces  here"

    def test_split_inside_hyperlink(self):
        doc = Document()
        para = doc.add_paragraph()
        para.add_run("See ")
        hyperlink = etree.SubElement(para._element, qn(NS_W, "hyperlink"))
        run = etree.SubElement(hyperlink, qn(NS_W, "r"))
        t = etree.SubElement(run, qn(NS_W, "t"))
        t.text = "the linked words"
        mgr = CommentManager(doc)
        cid = mgr.add_comment(
            para, "c", PersonInfo(author="A"), start_char=8, end_char=14
        )
        assert mgr.get_anchored_text(cid) == "linked"

    def test_br_counts_one_char(self):
        doc = Document()
        para = doc.add_paragraph()
        run = para.add_run("ab")
        etree.SubElement(run._element, qn(NS_W, "br"))
        para.add_run("cd")
        mgr = CommentManager(doc)
        cid = mgr.add_comment(
            para, "c", PersonInfo(author="A"), start_char=3, end_char=5
        )
        assert mgr.get_anchored_text(cid) == "cd"

    def test_validation_errors(self):
        doc, para, mgr = _setup()
        with pytest.raises(ValueError, match="not both"):
            mgr.add_comment(
                para, "c", PersonInfo(author="A"),
                start_run=0, end_run=0, start_char=0, end_char=3,
            )
        with pytest.raises(ValueError, match="together"):
            mgr.add_comment(para, "c", PersonInfo(author="A"), start_char=1)
        with pytest.raises(IndexError):
            mgr.add_comment(
                para, "c", PersonInfo(author="A"), start_char=0, end_char=999
            )
        with pytest.raises(ValueError):
            mgr.add_comment(
                para, "c", PersonInfo(author="A"), start_char=5, end_char=5
            )
        # No mutation on failure: no runs were split.
        assert len(para._element.findall(qn(NS_W, "r"))) == 1


class TestAddCommentOnText:
    def test_substring(self):
        doc, para, mgr = _setup()
        cid = mgr.add_comment_on_text(para, "brown", "c", PersonInfo(author="A"))
        assert mgr.get_anchored_text(cid) == "brown"

    def test_occurrence(self):
        doc, para, mgr = _setup("aba aba aba")
        cid = mgr.add_comment_on_text(
            para, "aba", "c", PersonInfo(author="A"), occurrence=2
        )
        assert mgr.get_anchored_text(cid) == "aba"
        # Splitting must not change the visible text.
        anchor = CommentAnchor(doc)
        assert anchor.paragraph_text(para._element) == "aba aba aba"
        # The SECOND occurrence (offset 4) was anchored: exactly one
        # commentRangeStart exists and the run text before it is "aba ".
        para_elem = para._element
        starts = para_elem.findall(qn(NS_W, "commentRangeStart"))
        assert len(starts) == 1
        children = list(para_elem)
        start_index = children.index(starts[0])
        preceding = "".join(
            t.text or ""
            for child in children[:start_index]
            if child.tag == qn(NS_W, "r")
            for t in child.findall(qn(NS_W, "t"))
        )
        assert preceding == "aba "

    def test_regex(self):
        doc, para, mgr = _setup()
        cid = mgr.add_comment_on_text(
            para, re.compile(r"qu\w+"), "c", PersonInfo(author="A")
        )
        assert mgr.get_anchored_text(cid) == "quick"

    def test_missing_match_raises_with_count(self):
        doc, para, mgr = _setup()
        with pytest.raises(ValueError, match="0 time"):
            mgr.add_comment_on_text(para, "zebra", "c", PersonInfo(author="A"))
        with pytest.raises(ValueError, match="1 time"):
            mgr.add_comment_on_text(
                para, "quick", "c", PersonInfo(author="A"), occurrence=2
            )
